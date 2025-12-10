# core/views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login
from django.contrib import messages
from django.views.decorators.csrf import ensure_csrf_cookie, csrf_exempt
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_GET
from django.http import JsonResponse, Http404
from collections import defaultdict
from django.db import transaction
from datetime import date, timedelta
from django.db.models import Sum, Q
from django.templatetags.static import static

import datetime as dt
from django.utils.timezone import make_aware
from datetime import datetime, time
from datetime import timedelta
from django.db.models import Sum
import re
import json
from rest_framework.response import Response
from rest_framework.views import APIView
from django.contrib.auth.models import User
from rest_framework.decorators import api_view, permission_classes, authentication_classes
from rest_framework.permissions import IsAuthenticated #neww
from rest_framework.authentication import SessionAuthentication
from rest_framework_simplejwt.authentication import JWTAuthentication
from django.template.loader import render_to_string
from django.utils import timezone
from rest_framework.decorators import api_view, permission_classes, authentication_classes, parser_classes

# import your models (Reservation was missing before)
from .models import UserBorrower, Item, Reservation, Feedback, DamageReport, BlockedDate, ReservationItem, TransactionCounter
from django.contrib.auth.hashers import check_password
from django.contrib.auth import update_session_auth_hash

# If you need DRF perms later:
# from rest_framework.decorators import api_view, permission_classes, authentication_classes
from rest_framework.permissions import IsAuthenticated, AllowAny

from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser

#notifications
import qrcode
from io import BytesIO
from django.core.files.base import ContentFile
from firebase_admin import messaging
from .models import DeviceToken, Notification
from .models import Notification
from .scheduler import run_scheduled_notifications


#stats
from django.db.models import Count
from django.db.models.functions import ExtractMonth
from django.utils.dateparse import parse_date
from django.views.decorators.http import require_GET
import csv
import io
from django.http import HttpResponse
import pandas as pd
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .models import AdminBorrow


#FORGOT PASSWORD
from django.core.mail import send_mail
from django.contrib import messages
from django.shortcuts import render, redirect
from django.conf import settings
import random

#SIGN UP
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes, force_str
from django.contrib.auth.tokens import default_token_generator
from django.core.mail import send_mail
from django.conf import settings

#logout
from django.shortcuts import redirect
from django.contrib.auth import logout as auth_logout

from .models import Notification, DeviceToken
from django.core.mail import EmailMultiAlternatives
from django.utils.html import strip_tags

from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from .models import Reservation, ReservationItem, UserBorrower, Item
import json

# --------------------------------------------
# PRIORITY SCORING FUNCTION (GLOBAL UTILITY)
# --------------------------------------------
def compute_priority(reservation):
    # ---- Event Type Weight (from priority_detail) ----
    detail = (reservation.priority_detail or "").lower()

    if "funeral" in detail:
        event_weight = 5
    elif "emergency" in detail or "barangay" in detail:
        event_weight = 4
    elif "inside" in detail:
        event_weight = 3
    elif "outside" in detail:
        event_weight = 2
    else:
        event_weight = 1

    # ---- Borrower Behavior Score ----
    borrower = reservation.userborrower
    behavior_score = getattr(borrower, "behavior_score", 0)

    # ---- Request Timing Weight ----
    request_date = reservation.created_at.date()
    borrow_date = reservation.date_borrowed
    diff_days = (borrow_date - request_date).days

    if diff_days >= 7:
        timing_weight = 3
    elif diff_days >= 3:
        timing_weight = 2
    else:
        timing_weight = 1

    # ---- Urgency Weight (fixed logic) ----
    urgency_weight = 1
    for r_item in reservation.items.all():

        reserved = ReservationItem.objects.filter(
            item=r_item.item,
            reservation__status__in=['approved', 'in use']
        ).aggregate(total=Sum("quantity"))["total"] or 0

        available = r_item.item.qty - reserved

        if available <= 1:   # nearly out of stock
            urgency_weight = 2

    # ---- Final priority formula ----
    priority = (
        (event_weight * 0.40) +
        (behavior_score * 0.30) +
        (timing_weight * 0.20) +
        (urgency_weight * 0.10)
    )

    return round(priority, 2)


# -----------------------
# Web views (templates)
# -----------------------

def admin_login(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")

        user = authenticate(request, username=username, password=password)

        if user is not None:
            if user.is_staff:  # Only allow admin/staff
                login(request, user)
                return redirect('dashboard')  # Dashboard URL (to be created)
            else:
                messages.error(request, "You do not have admin access")
        else:
            messages.error(request, "Invalid username or password")

    return render(request, "login.html")


@login_required
def dashboard(request):
    # Summary cards
    total_users = UserBorrower.objects.count()
    total_items = Item.objects.count()
    total_transactions = Reservation.objects.count()
    total_borrowed = Reservation.objects.filter(status__iexact='in use').count()

    # PIE - Item Category Distribution
    category_data = Item.objects.values("category").annotate(count=Count("item_id"))
    pie_labels = [c["category"] for c in category_data]
    pie_values = [c["count"] for c in category_data]

    # BAR - Monthly Transactions (based on date_borrowed)
    current_year = timezone.now().year
    monthly_data = (
        Reservation.objects.filter(date_borrowed__year=current_year)
        .annotate(month=ExtractMonth("date_borrowed"))
        .values("month")
        .annotate(count=Count("id"))
        .order_by("month")
    )

    month_names = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    month_counts = {m["month"]: m["count"] for m in monthly_data}
    bar_labels = month_names
    bar_values = [month_counts.get(i, 0) for i in range(1, 13)]

    # DONUT - Borrowed vs Returned
    borrowed = Reservation.objects.filter(status__iexact='in use').count()
    returned = Reservation.objects.filter(status__iexact='returned').count()
    total = borrowed + returned
    borrowed_percent = round((borrowed / total * 100), 1) if total else 0
    returned_percent = round((returned / total * 100), 1) if total else 0
    borrowed_vs_returned = {
        "borrowed": borrowed_percent,
        "returned": returned_percent,
    }

    context = {
        "total_users": total_users,
        "total_items": total_items,
        "total_transactions": total_transactions,
        "total_borrowed": total_borrowed,
        "pie_labels": pie_labels,
        "pie_values": pie_values,
        "bar_labels": bar_labels,
        "bar_values": bar_values,
        "borrowed_vs_returned": borrowed_vs_returned,
    }
    return render(request, "dashboard.html", context)

#NEW
def forgot_password(request):
    show_code_container = False
    email_value = ""  # store email to keep it in the input

    if request.method == 'POST':
        # When admin clicks "Send Reset Code"
        if 'send_code' in request.POST:
            email = request.POST.get('email')
            email_value = email  # keep value for re-render

            try:
                user = User.objects.get(email=email)
                code = random.randint(100000, 999999)
                request.session['reset_email'] = email
                request.session['reset_code'] = str(code)

                send_mail(
                    subject="TrailLend Password Reset Code",
                    message=f"Your password reset code is {code}. Please use this code to verify your identity.",
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[email],
                    fail_silently=False,
                )

                messages.success(request, "A reset code has been sent to your email. Please check your inbox.")
                show_code_container = True

            except User.DoesNotExist:
                messages.error(request, "This email doesn't exist.")

        # When admin clicks "Verify Code"
        elif 'verify_code' in request.POST:
            input_code = request.POST.get('reset_code')
            session_code = request.session.get('reset_code')
            email_value = request.session.get('reset_email', '')

            if input_code == session_code:
                messages.success(request, "Code verified successfully! You can now reset your password.")
                return redirect('verify_reset_code')
            else:
                messages.error(request, "Invalid or incorrect code.")
                show_code_container = True

        # When admin clicks "Resend Code"
        elif 'resend_code' in request.POST:
            email = request.session.get('reset_email')
            email_value = email
            if email:
                code = random.randint(100000, 999999)
                request.session['reset_code'] = str(code)
                send_mail(
                    subject="TrailLend Password Reset Code (Resent)",
                    message=f"Your new password reset code is {code}.",
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[email],
                    fail_silently=False,
                )
                messages.success(request, "A new code has been sent to your email.")
                show_code_container = True
            else:
                messages.error(request, "No email session found. Please enter your email again.")

    return render(request, "forgot_password.html", {
        'show_code_container': show_code_container,
        'email': email_value or request.session.get('reset_email', '')
    })

#NEW
def verify_reset_code(request):
    """
    Page for entering a new password after verifying the reset code.
    """
    email = request.session.get('reset_email')

    if not email:
        messages.error(request, "Session expired. Please enter your email again.")
        return redirect('forgot_password')

    if request.method == 'POST':
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')

        if not new_password or not confirm_password:
            messages.error(request, "Please fill in all fields.")
        elif new_password != confirm_password:
            messages.error(request, "Passwords do not match.")
        else:
            try:
                user = User.objects.get(email=email)
                user.set_password(new_password)
                user.save()

                # Clear session data
                request.session.pop('reset_email', None)
                request.session.pop('reset_code', None)

                # Show success message (used by popup)
                messages.success(request, "Your password has been successfully changed.")
                return render(request, "verify_reset_code.html")

            except User.DoesNotExist:
                messages.error(request, "User not found. Please try again.")

    return render(request, "verify_reset_code.html", {"email": email})


def inventory(request):
    items = Item.objects.all()

    # Get filter parameters from GET request
    category = request.GET.get('category', '')
    status = request.GET.get('status', '')
    sort = request.GET.get('sort', 'newest')

    # Apply category filter
    if category:
        items = items.filter(category=category)

    # Apply status filter
    if status:
        items = items.filter(status=status)

    # Apply sorting
    if sort == 'newest':
        items = items.order_by('-item_id')
    else:
        items = items.order_by('item_id')

    total_items = items.count()

    # ✅ NEW: Total quantity of all items
    total_quantity = Item.objects.aggregate(total=Sum('qty'))['total'] or 0

    context = {
        'items': items,
        'category': category,
        'status': status,
        'sort': sort,
        'total_items': total_items,
        'total_quantity': total_quantity,   # ← ADDED
    }

    return render(request, 'inventory.html', context)

def inventory_createitem(request):
    if request.method == "POST":
        name = request.POST.get('item_name')
        qty = request.POST.get('quantity')
        category = request.POST.get('category')
        description = request.POST.get('description', '')
        image = request.FILES.get('item_image')
        status = request.POST.get('item_status', 'Available')
        owner = request.POST.get('item_owner', 'Barangay Kauswagan')

        # Create the item using correct model field names
        Item.objects.create(
            name=name,
            qty=qty,
            category=category,
            description=description,
            image=image,
            owner=owner,
            status=status,
        )

        return redirect('inventory')

    return render(request, "inventory_createitem.html")

@csrf_exempt
def run_scheduler_api(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST only"}, status=405)

    sent = run_scheduled_notifications()
    return JsonResponse({"sent": sent})

def inventory_edit(request, item_id):
    item = Item.objects.get(item_id=item_id)

    if request.method == 'POST':
        item.name = request.POST.get('name')
        item.qty = request.POST.get('qty')
        item.description = request.POST.get('description')
        item.category = request.POST.get('category')
        item.status = request.POST.get('status')
        item.owner = request.POST.get('owner')

        if 'image' in request.FILES:
            item.image = request.FILES['image']

        item.save()
        return redirect('inventory')

    return render(request, "inventory_edit.html", {'item': item})


def inventory_detail(request, item_id):
    item = Item.objects.get(item_id=item_id)
    return render(request, "inventory_detail.html", {'item': item})


def inventory_delete(request):
    return render(request, "inventory_confirm_delete.html")



def verification(request):
    return render(request, 'verification.html')

def ensure_datetime(value):
    if not value:
        return None
    if isinstance(value, datetime):
        return value
    return datetime.combine(value, time.min)  # convert date to datetime 00:00


def transaction_log(request):

    def fmt(dt):
        """Format datetime safely for display."""
        if not dt:
            return None
        try:
            # If time = 00:00:00 → treat as date only
            if dt.hour == 0 and dt.minute == 0 and dt.second == 0:
                return dt.date()
        except:
            pass
        return dt

    transactions = []

    # ===============================
    # 1. USER RESERVATIONS
    # ===============================
    qs = (
        Reservation.objects
        .select_related("userborrower")
        .prefetch_related("items__item")
        .order_by("-id")
    )

    for r in qs:
        borrower_name = getattr(r.userborrower, "full_name", "Unknown")

        # MULTIPLE ITEMS
        item_list = [ri.item_name for ri in r.items.all()]
        qty_list = [ri.quantity for ri in r.items.all()]

        # CONTACT FALLBACK
        contact = (r.contact or "").strip()
        if not contact or contact.lower() == "n/a":
            contact = getattr(r.userborrower, "contact_number", "N/A")

        transactions.append({
            "transaction_id": r.transaction_id,
            "user_name": borrower_name,
            "item_list": item_list,
            "qty_list": qty_list,
            "contact": contact,
            "created_at": r.created_at,
            "date_receive": fmt(r.date_receive),
            "date_returned": fmt(r.date_returned),
            "delivered_by": r.delivered_by,
            "status": r.status,
        })

    # ===============================
    # 2. ADMIN DIRECT BORROW
    # ===============================
    ab_list = AdminBorrow.objects.select_related("item").order_by("-id")

    for ab in ab_list:
        transactions.append({
            "transaction_id": ab.transaction_id,
            "user_name": ab.borrower_name,
            "item_list": [ab.item.name],  # ADMIN = single item
            "qty_list": [ab.quantity],
            "contact": ab.contact_number,
            "created_at": ab.created_at,
            "date_receive": fmt(ab.date),
            "date_returned": fmt(ab.return_date if ab.status == "Returned" else None),
            "delivered_by": ab.delivered_by,
            "status": ab.status,
        })

    # ===============================
    # 4. SORT ALL COMBINED TRANSACTIONS
    # ===============================
    transactions = sorted(
        transactions,
        key=lambda x: x["created_at"],
        reverse=True
    )

    return render(request, "transaction_history.html", {"transactions": transactions})



# Statistics


@login_required
def statistics(request):
    all_items = Item.objects.all()

    context = {
        "items": all_items,
    }
    return render(request, "statistics.html", context)


def statistics_data(request):

    # Filters from request
    start = request.GET.get("start")
    end = request.GET.get("end")
    status_filter = request.GET.get("status", "all")
    category_filter = request.GET.get("category", "all")
    report_type_filter = request.GET.get("report_type", "all")

    # Base: get all reservations
    reservations = (
        Reservation.objects
        .select_related("userborrower")
        .prefetch_related("items__item", "damage_reports")
        .all()
    )

    # Apply DATE filters (borrowed date)
    if start:
        reservations = reservations.filter(date_borrowed__gte=parse_date(start))
    if end:
        reservations = reservations.filter(date_borrowed__lte=parse_date(end))

    # Apply STATUS filter
    if status_filter != "all":
        reservations = reservations.filter(status=status_filter)

    results = []

    # Build results
    for r in reservations:

        # Determine report type (Damage / Loss / None)
        report = r.damage_reports.first()
        report_type = report.report_type.lower() if report else "none"

        # Filter by report type
        if report_type_filter != "all":
            if report_type_filter != report_type:
                continue

        # Loop through all items in reservation
        for ri in r.items.all():

            # Category filter
            if category_filter != "all" and ri.item.category != category_filter:
                continue

            results.append({
                "item_name": ri.item.name,
                "category": ri.item.category,
                "borrower_name": r.userborrower.full_name if r.userborrower else "Unknown",
                "borrowed_at": r.date_borrowed.strftime("%Y-%m-%d"),
                "returned_at": r.date_return.strftime("%Y-%m-%d"),
                "report_type": report_type.capitalize(),
                "status": r.status,
            })

    return JsonResponse({"transactions": results})


@login_required
def export_excel(request):
    start_date = request.GET.get("start")
    end_date = request.GET.get("end")
    status_filter = request.GET.get("status")
    category_filter = request.GET.get("category")
    report_filter = request.GET.get("report_type")

    qs = (
        Reservation.objects
        .select_related("userborrower")
        .prefetch_related("items__item", "damage_reports")
        .all()
    )

    # Filters
    if start_date:
        qs = qs.filter(date_borrowed__gte=parse_date(start_date))
    if end_date:
        qs = qs.filter(date_borrowed__lte=parse_date(end_date))
    if status_filter and status_filter != "all":
        qs = qs.filter(status__iexact=status_filter)

    data = []

    for r in qs:
        report = r.damage_reports.first()
        report_type = report.report_type if report else "None"

        for ri in r.items.all():

            # Category filter
            if category_filter != "all" and ri.item.category != category_filter:
                continue

            # Damage/Loss filter
            if report_filter == "damage" and report_type != "Damage":
                continue
            if report_filter == "loss" and report_type != "Loss":
                continue

            data.append({
                "Item Name": ri.item.name,
                "Category": ri.item.category,
                "Borrower": r.userborrower.full_name,
                "Borrowed At": r.date_borrowed.strftime("%Y-%m-%d"),
                "Returned At": r.date_return.strftime("%Y-%m-%d") if r.date_return else "",
                "Report Type": report_type,
                "Status": r.status.capitalize(),
            })

    df = pd.DataFrame(data)
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Report")

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = "attachment; filename=report.xlsx"
    return response


@login_required
def export_pdf(request):
    start_date = request.GET.get("start")
    end_date = request.GET.get("end")
    status_filter = request.GET.get("status")
    category_filter = request.GET.get("category")
    report_filter = request.GET.get("report_type")

    qs = (
        Reservation.objects
        .select_related("userborrower")
        .prefetch_related("items__item", "damage_reports")
        .all()
    )

    if start_date:
        qs = qs.filter(date_borrowed__gte=parse_date(start_date))
    if end_date:
        qs = qs.filter(date_borrowed__lte=parse_date(end_date))
    if status_filter and status_filter != "all":
        qs = qs.filter(status__iexact=status_filter)

    transactions = []

    for r in qs:
        report = r.damage_reports.first()
        report_type = report.report_type if report else "None"

        for ri in r.items.all():

            if category_filter != "all" and ri.item.category != category_filter:
                continue

            if report_filter == "damage" and report_type != "Damage":
                continue
            if report_filter == "loss" and report_type != "Loss":
                continue

            transactions.append({
                "item_name": ri.item.name,
                "category": ri.item.category,
                "borrower_name": r.userborrower.full_name,
                "borrowed_at": r.date_borrowed.strftime("%Y-%m-%d"),
                "returned_at": r.date_return.strftime("%Y-%m-%d") if r.date_return else "",
                "report_type": report_type,
                "status": r.status.capitalize(),
            })

    logo_path = request.build_absolute_uri(static("barangay_kauswagan_logo.png"))

    html = render_to_string(
        "pdf_template.html",
        {"transactions": transactions, "logo_path": logo_path}
    )

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = "attachment; filename=report.pdf"

    pisa.CreatePDF(html, dest=response)
    return response

@login_required
def export_docx(request):
    import os
    import io
    from django.http import HttpResponse
    from django.conf import settings
    from django.utils.dateparse import parse_date
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # -------------------------
    # SHADING (HEADER BG COLOR)
    # -------------------------
    def shade(cell, color="E8EEF3"):
        """Safe shading setter."""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    # -------------------------
    # FILTERS
    # -------------------------
    start_date = request.GET.get("start")
    end_date = request.GET.get("end")
    status_filter = request.GET.get("status")
    category_filter = request.GET.get("category")
    report_filter = request.GET.get("report_type")

    qs = (
        Reservation.objects
        .select_related("userborrower")
        .prefetch_related("items__item", "damage_reports")
        .all()
    )

    if start_date:
        qs = qs.filter(date_borrowed__gte=parse_date(start_date))
    if end_date:
        qs = qs.filter(date_borrowed__lte=parse_date(end_date))
    if status_filter and status_filter != "all":
        qs = qs.filter(status__iexact=status_filter)

    # -------------------------
    # BUILD TRANSACTIONS
    # -------------------------
    transactions = []
    for r in qs:
        rep = r.damage_reports.first()
        rep_type = rep.report_type if rep else "None"

        for ri in r.items.all():

            if category_filter != "all" and ri.item.category != category_filter:
                continue

            if report_filter == "damage" and rep_type != "Damage":
                continue
            if report_filter == "loss" and rep_type != "Loss":
                continue

            transactions.append({
                "item_name": ri.item.name,
                "category": ri.item.category,
                "borrower_name": r.userborrower.full_name,
                "borrowed_at": r.date_borrowed.strftime("%Y-%m-%d"),
                "returned_at": r.date_return.strftime("%Y-%m-%d") if r.date_return else "",
                "report_type": rep_type,
                "status": r.status.capitalize(),
            })

    # -------------------------
    # CREATE DOCX
    # -------------------------
    doc = Document()

    # PAGE SETUP LANDSCAPE
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # -------------------------
    # HEADER TABLE (LOGO + CENTER TEXT)
    # -------------------------
    header = doc.add_table(rows=1, cols=3)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = header.rows[0].cells

    # LEFT LOGO
    left = hdr_cells[0]
    left_par = left.paragraphs[0]
    left_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        logo_path = os.path.join(settings.BASE_DIR, "core", "static", "barangay_kauswagan_logo.png")
        run = left_par.add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    except:
        left_par.add_run("")

    # CENTER TEXT
    center = hdr_cells[1]
    c = center.paragraphs[0]
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    texts = [
        ("REPUBLIC OF THE PHILIPPINES", True),
        ("City of Cagayan de Oro", False),
        ("BARANGAY KAUSWAGAN", False),
        ("OFFICE OF THE PUNONG BARANGAY", True),
        ("FACILITIES AND PROPERTIES", True)
    ]

    for text, bold in texts:
        r = c.add_run(text + "\n")
        r.bold = bold
        r.font.size = Pt(12)
        if "FACILITIES" in text:
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0, 51, 102)

    # RIGHT CELL (EMPTY — same as PDF)
    hdr_cells[2].text = ""

    doc.add_paragraph("")  # spacing
    doc.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # TITLE
    title = doc.add_paragraph("Item Activity Report")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph("")

    # -------------------------
    # TABLE SETUP (SAME AS PDF)
    # -------------------------
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Item Name", "Category", "Borrower",
        "Borrowed At", "Returned At",
        "Report Type", "Status"
    ]

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        shade(hdr[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    # -------------------------
    # TABLE ROWS
    # -------------------------
    for t in transactions:
        row = table.add_row().cells
        row[0].text = t["item_name"]
        row[1].text = t["category"]
        row[2].text = t["borrower_name"]
        row[3].text = t["borrowed_at"]
        row[4].text = t["returned_at"]
        row[5].text = t["report_type"]
        row[6].text = t["status"]

        for cell in row:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # -------------------------
    # EXPORT DOCX FILE
    # -------------------------
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = "attachment; filename=report.docx"
    return response



def set_cell_shading(cell, fill_color):
    """
    Applies background shading to a cell in python-docx.
    fill_color = hex color like 'E8EEF3'
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def change_pass(request):
    return render(request, 'change-password.html')


def list_of_users(request):
    profiles = UserBorrower.objects.select_related('user').all()
    return render(request, 'list_of_users.html', {'profiles': profiles})


def logout(request):
    auth_logout(request)
    return redirect('login')


# -----------------------
# API views (JSON)
# -----------------------
@csrf_exempt
def api_register(request):
    """
    API endpoint for registration with HTML email and local dev IP support.
    """
    if request.method == "POST":
        try:
            data = json.loads(request.body or "{}")
            username = data.get("username")
            password = data.get("password")
            confirm_password = data.get("confirmPassword")
            full_name = data.get("name")
            contact_number = data.get("contactNumber")
            address = data.get("address")
            email = data.get("email")


            # Validation
            if not all([username, password, confirm_password, full_name, email]):
                return JsonResponse({"success": False, "message": "Missing required fields"}, status=400)


            if User.objects.filter(username=username).exists():
                return JsonResponse({"success": False, "message": "Username already exists"}, status=400)


            if User.objects.filter(email=email).exists():
                return JsonResponse({"success": False, "message": "Email already registered"}, status=400)


            if password != confirm_password:
                return JsonResponse({"success": False, "message": "Passwords do not match"}, status=400)


            # Create inactive user
            user = User.objects.create_user(
                username=username,
                password=password,
                email=email,
                is_active=False
            )


            UserBorrower.objects.create(
                user=user,
                full_name=full_name,
                contact_number=contact_number,
                address=address
            )


            # Generate verification link (use your local IP for now)
            uid = urlsafe_base64_encode(force_bytes(user.pk))
            token = default_token_generator.make_token(user)


            # 🧩 Local development link — works, but user won't see the IP
            verify_url = f"http://10.196.129.244:8000/api/verify-email/{uid}/{token}/"


            # HTML Email Template
            html_message = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
              <meta charset="UTF-8" />
              <meta name="viewport" content="width=device-width, initial-scale=1.0" />
              <title>Verify Your Email - TrailLend</title>
            </head>
            <body style="font-family:'Poppins',Arial,sans-serif; background-color:#f4f6f9; padding:30px;">
              <table align="center" style="max-width:600px; background:#fff; border-radius:12px; overflow:hidden; box-shadow:0 4px 12px rgba(0,0,0,0.1);">
                <tr>
                  <td style="background-color:#1976D2; text-align:center; padding:30px;">
                    <img src="https://i.postimg.cc/Dw8pYLL5/TRAILLEND-ICON.png" alt="TrailLend Logo" width="80" />
                    <h1 style="color:#fff; margin:10px 0 0;">TrailLend</h1>
                    <p style="color:#cce6ff;">Empowering the Community Together 🌿</p>
                  </td>
                </tr>
                <tr>
                  <td style="padding:30px;">
                    <h2 style="color:#1976D2;">Email Verification Required</h2>
                    <p style="color:#333;">Hi {full_name},</p>
                    <p style="color:#555;">Thank you for registering on <strong>TrailLend</strong>!
                    To activate your account, please verify your email address by clicking the button below:</p>


                    <div style="text-align:center; margin:30px 0;">
                      <a href="{verify_url}"
                         style="background-color:#1976D2; color:#fff; text-decoration:none;
                                padding:14px 28px; border-radius:8px; font-weight:bold; display:inline-block;">
                        Verify My Email
                      </a>
                    </div>


                    <p style="color:#777; font-size:14px;">If you didn’t create this account, please ignore this email.</p>
                    <hr style="border:none; border-top:1px solid #eee; margin:30px 0;">
                    <p style="color:#999; font-size:12px; text-align:center;">
                      © 2025 TrailLend • Barangay General Services Office<br>
                      Please do not reply directly to this message.
                    </p>
                  </td>
                </tr>
              </table>
            </body>
            </html>
            """


            send_mail(
                subject="Verify Your TrailLend Account",
                message="Please verify your TrailLend account.",
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[email],
                fail_silently=False,
                html_message=html_message
            )


            return JsonResponse({
                "success": True,
                "message": "Registration successful! Check your email to verify your account."
            }, status=201)


        except Exception as e:
            import traceback
            traceback.print_exc()
            return JsonResponse({"success": False, "message": str(e)}, status=400)


    return JsonResponse({"success": False, "message": "Invalid request method"}, status=405)




# Verify Email Endpoint
@api_view(["GET"])
@permission_classes([AllowAny])
def verify_email(request, uidb64, token):
    """
    Activates the user's account and hides IP on success.
    """
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        user = User.objects.get(pk=uid)


        if default_token_generator.check_token(user, token):
            user.is_active = True
            user.save()


            # Deep link to app
            deep_link = "com.traillend.app://verified"


            # Beautiful success page (no visible IP)
            html = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
              <meta charset="UTF-8">
              <meta name="viewport" content="width=device-width, initial-scale=1.0">
              <title>Email Verified - TrailLend</title>
              <meta http-equiv="refresh" content="3;url={deep_link}">
              <style>
                body {{
                    font-family: 'Poppins', Arial, sans-serif;
                    background-color: #f4f6f9;
                    text-align: center;
                    padding: 60px;
                }}
                .card {{
                    background: #fff;
                    border-radius: 16px;
                    max-width: 500px;
                    margin: auto;
                    padding: 40px;
                    box-shadow: 0 4px 10px rgba(0,0,0,0.1);
                }}
                .btn {{
                    background-color: #1976D2;
                    color: #fff;
                    text-decoration: none;
                    padding: 12px 24px;
                    border-radius: 8px;
                    display: inline-block;
                    margin-top: 20px;
                    font-weight: bold;
                }}
                .btn:hover {{
                    background-color: #145CA8;
                }}
              </style>
            </head>
            <body>
              <div class="card">
                <img src="https://i.postimg.cc/Dw8pYLL5/TRAILLEND-ICON.png" width="90" />
                <h1 style="color:#1976D2;">Email Verified!</h1>
                <p style="color:#333;">Your TrailLend account has been successfully verified.</p>
                <p style="color:#555;">You can now log in to your account.</p>
                <a href="{deep_link}" class="btn">Open TrailLend App</a>
              </div>
            </body>
            </html>
            """


            return HttpResponse(html)


        else:
            return JsonResponse({"success": False, "message": "Invalid or expired verification link."}, status=400)
    except Exception as e:
        return JsonResponse({"success": False, "message": str(e)}, status=400)



@csrf_exempt
def api_login(request):
    """
    Mobile login API — allows restricted borrowers to login,
    but returns their borrower_status so frontend can show modal.
    """
    if request.method == "POST":
        try:
            data = json.loads(request.body or "{}")
            username = data.get("username")
            password = data.get("password")

            if not username or not password:
                return JsonResponse({
                    "success": False,
                    "message": "Username and password required"
                }, status=400)

            user = authenticate(request, username=username, password=password)

            if user is None:
                return JsonResponse({
                    "success": False,
                    "message": "Invalid credentials"
                }, status=401)

            borrower = UserBorrower.objects.get(user=user)

            # ⭐ ALLOW LOGIN FOR RESTRICTED ACCOUNTS
            # DO NOT block login with 403
            # Instead, return borrower_status to app
            refresh = RefreshToken.for_user(user)

            return JsonResponse({
                "success": True,
                "message": "Login successful",
                "refresh": str(refresh),
                "access": str(refresh.access_token),
                "borrower_status": borrower.borrower_status,   # ⭐ VERY IMPORTANT
                "late_count": borrower.late_count              # ⭐ Add for UI if needed
            }, status=200)

        except Exception as e:
            return JsonResponse({
                "success": False,
                "message": str(e)
            }, status=400)

    return JsonResponse({
        "success": False,
        "message": "Invalid request method"
    }, status=405)




@api_view(['GET'])
@permission_classes([AllowAny])
def pending_requests_api(request):
    qs = (
    Reservation.objects
    .filter(status='pending')
    .select_related('userborrower')
    .prefetch_related('items__item')
    .order_by('-priority_score', '-created_at')   # IMPORTANT
)


    html = render_to_string(
        'pending_requests_list.html',
        {'pending_request': qs},
        request=request
    )
    return Response({
        'html': html,
        'count': qs.count()
    })


@api_view(['GET'])
@authentication_classes([SessionAuthentication, JWTAuthentication])
@permission_classes([IsAuthenticated])
def reservation_detail_api(request, pk: int):

    r = get_object_or_404(
        Reservation.objects
        .select_related('userborrower')
        .prefetch_related('items__item'),
        pk=pk
    )

    # -------------------------------
    # FIXED: borrower structured object
    # -------------------------------
    borrower_data = {
        "full_name": r.userborrower.full_name if r.userborrower else "Unknown",
        "contact_number": (
            r.userborrower.contact_number 
            if r.userborrower and hasattr(r.userborrower, "contact_number")
            else ""
        )
    }

    # -------------------------------
    # FIXED: items list (correct item name)
    # -------------------------------
    items = []
    for it in r.items.all():
        items.append({
            "item_name": it.item.name if it.item else it.item_name,
            "quantity": it.quantity,
            "image": request.build_absolute_uri(it.item.image.url)
                     if it.item and it.item.image else ""
        })
        latest_decline = Notification.objects.filter(
            reservation=r,
            type="rejection"
        ).order_by("-created_at").first()

        decline_reason = latest_decline.reason if latest_decline else ""

    data = {
        "id": r.id,
        "transaction_id": r.transaction_id,

        # FIXED: send borrower as object
        "userborrower": borrower_data,
        "contact_number": borrower_data["contact_number"],

        "date_borrowed": r.date_borrowed.strftime('%Y-%m-%d'),
        "date_return": r.date_return.strftime('%Y-%m-%d'),
        "priority_display": r.priority,   # JS expects "priority_display"
        "message": r.message or "No message provided",
        "status": r.status,
        "reason": decline_reason,
        
        "letter_image": request.build_absolute_uri(r.letter_image.url) 
                        if r.letter_image else "",
        "valid_id_image": request.build_absolute_uri(r.valid_id_image.url) 
                        if r.valid_id_image else "",
        "items": items
    }

    return Response(data)
    

# --------------------------------------------
# PRIORITY CONFLICT DETECTION FUNCTION
# --------------------------------------------
def detect_conflicts_and_notify(reservation):
    main_start = reservation.date_borrowed
    main_end = reservation.date_return

    for r_item in reservation.items.all():

        item = r_item.item
        approved_qty = r_item.quantity
        total_stock = item.qty

        # Remaining stock AFTER the approved reservation consumes quantity
        remaining_qty = total_stock - approved_qty

        # Find overlapping pending reservations for the same item + dates
        overlapping = ReservationItem.objects.filter(
            item=item,
            reservation__status='pending',
            reservation__date_borrowed__lte=main_end,
            reservation__date_return__gte=main_start
        ).exclude(reservation=reservation)

        # Recommended dates (computed once)
        alt_dates = find_alternative_dates(item, main_start, main_end)

        for other in overlapping:
            other_res = other.reservation
            borrower = other_res.userborrower
            requested_qty = other.quantity

            # -------------------------------------
            # ⭐ ONLY notify if requested exceeds remaining
            # -------------------------------------
            if requested_qty > remaining_qty:

                Notification.objects.create(
                    user=borrower,
                    reservation=other_res,
                    title="Date Conflict Detected",
                    message=(
                        f"Your reservation for {main_start} → {main_end} cannot be accommodated "
                        f"because the item stock is insufficient.\n\n"
                        f"Requested: {requested_qty}, Available: {remaining_qty}\n\n"
                        f"Recommended available dates:\n"
                        f"{', '.join(alt_dates) if alt_dates else 'No alternative dates available'}"
                    ),
                    type="conflict_alert",
                    is_sent=True
                )



def find_alternative_dates(item, start, end, limit=5):
    suggestions = []

    check_start = end + timedelta(days=1)
    day = check_start

    for _ in range(60):  # check 60 future days
        overlap = ReservationItem.objects.filter(
            item=item,
            reservation__status__in=['pending', 'approved', 'in use'],
            reservation__date_borrowed__lte=day,
            reservation__date_return__gte=day
        ).aggregate(total=Sum("quantity"))["total"] or 0

        if overlap < item.qty:
            suggestions.append(day.isoformat())

            if len(suggestions) >= limit:
                break

        day += timedelta(days=1)

    return suggestions




@api_view(['POST'])
@authentication_classes([SessionAuthentication, JWTAuthentication])
@permission_classes([IsAuthenticated])
@transaction.atomic
def reservation_update_api(request, pk: int):

    r = get_object_or_404(
        Reservation.objects.prefetch_related('items__item').select_related('userborrower'),
        pk=pk
    )

    new_status = request.data.get('status')
    reason_text = (request.data.get('reason') or '').strip()


    allowed = {'approved', 'declined', 'in use', 'returned', 'pending', 'cancelled'}
    if new_status not in allowed:
        return Response({'status': 'error', 'message': 'Invalid status'}, status=400)

    # =====================================================
    # ❌ STOCK HANDLING — REMOVED (DATE-BASED SYSTEM)
    # =====================================================
    # No deduct_stock()
    # No restore_stock()

    prev = r.status

    # =====================================================
    # TIMESTAMPS
    # =====================================================
    if new_status == "approved":
        r.approved_at = timezone.now()

            # ⭐ Automatically detect conflicts + notify borrowers
        detect_conflicts_and_notify(r)

    elif new_status == "in use":
        r.date_receive = timezone.now()

    elif new_status == "returned":
        r.date_returned = timezone.now()

    r.status = new_status
    r.save()

    # =====================================================
    # SMART REMINDER SCHEDULER (Correct logic + no instant reminders)
    # =====================================================
    from datetime import datetime, time, timedelta
    from django.utils.timezone import make_aware
    now = timezone.now()

    def push_future(dt):
        if dt <= now:
            return now + timedelta(minutes=2)
        return dt

    def schedule_smart_alerts(reservation):
        borrower = reservation.userborrower

        # ---------- 1️ Return Reminder
        return_day_before = reservation.date_return - timedelta(days=1)
        dt1 = make_aware(datetime.combine(return_day_before, time(18, 0)))
        dt1 = push_future(dt1)

        Notification.objects.create(
            user=borrower,
            reservation=reservation,
            title="Return Reminder",
            message="Your borrowed items are due tomorrow. Please be ready to return.",
            type="return_reminder",
            scheduled_at=dt1,
            is_sent=False
        )

        # ---------- 2️ Claim Reminder
        dt2 = make_aware(datetime.combine(reservation.date_borrowed, time(6, 0)))
        dt2 = push_future(dt2)

        Notification.objects.create(
            user=borrower,
            reservation=reservation,
            title="Claim Reminder",
            message="You may now claim your reserved items today.",
            type="claim_reminder",
            scheduled_at=dt2,
            is_sent=False
        )

        # ---------- 3️ Claim Delay Warning
        dt3 = reservation.created_at + timedelta(hours=1)
        dt3 = push_future(dt3)

        Notification.objects.create(
            user=borrower,
            reservation=reservation,
            title="Claim Delay Warning",
            message="You haven't claimed your items yet. Please claim them as soon as possible.",
            type="warning_claim_delay",
            scheduled_at=dt3,
            is_sent=False
        )

    # Run scheduler ONLY when approved
    if new_status == "approved":
        schedule_smart_alerts(r)

    # =====================================================
    # INSTANT NOTIFICATIONS
    # =====================================================

    if new_status == "pending":
        Notification.objects.create(
            user=r.userborrower,
            reservation=r,
            title="Pending Reservation",
            message=f"Your reservation ({r.transaction_id}) is pending approval.",
            type="pending",
            is_sent=True
        )

    elif new_status == "approved":

        qr_lines = [
            f"Transaction: {r.transaction_id}",
            f"Borrower: {r.userborrower.full_name}",
            f"Borrowed: {r.date_borrowed} → {r.date_return}",
            "",
            "Items:"
        ]
        for it in r.items.all():
            qr_lines.append(f" • {it.item_name} x{it.quantity}")

        qr_data = "\n".join(qr_lines)

        qr_img = qrcode.make(qr_data)
        buffer = BytesIO()
        qr_img.save(buffer, format='PNG')
        qr_file = ContentFile(buffer.getvalue(), f"qr_{r.transaction_id}.png")

        notif = Notification.objects.create(
            user=r.userborrower,
            reservation=r,
            title="Reservation Approved",
            message="Your reservation has been approved! Your QR code is ready.",
            type="approval",
            is_sent=True
        )
        notif.qr_code.save(f"qr_{r.transaction_id}.png", qr_file)

    elif new_status == "declined":
        Notification.objects.create(
            user=r.userborrower,
            reservation=r,
            title="Reservation Declined",
            message="Your reservation was declined.",
            reason=reason_text,
            type="rejection",
            is_sent=True
        )

    return Response({"status": "success"})





PRIORITY_LABELS = {
    "High":   "High — Bereavement",
    "Low":    "Low — General",
}



def schedule_smart_notifications(reservation):
    borrower = reservation.userborrower
    item = reservation.item

    claim_date = reservation.date_borrowed
    return_date = reservation.date_return

    # --- 1. Reminder @ 6PM day before return date ---
    reminder1_time = datetime.combine(return_date - timedelta(days=1), time(18, 0))

    Notification.objects.create(
        user=borrower,
        reservation=reservation,
        title="Return Reminder",
        message=f"Your borrowed item '{item.name}' is due tomorrow ({return_date}).",
        type="return_reminder_smart",
        scheduled_at=reminder1_time,
    )

    # --- 2. Reminder @ 6AM on claim day ---
    reminder2_time = datetime.combine(claim_date, time(6, 0))

    Notification.objects.create(
        user=borrower,
        reservation=reservation,
        title="Claim Reminder",
        message=f"Your item '{item.name}' can be claimed today.",
        type="claim_reminder_smart",
        scheduled_at=reminder2_time,
    )

    # --- 3. Reminder @ 1 hour AFTER claiming (only if claimed) ---
    if reservation.date_receive:
        reminder3_time = reservation.date_receive + timedelta(hours=1)
        Notification.objects.create(
            user=borrower,
            reservation=reservation,
            title="Follow-Up Reminder",
            message=f"It has been 1 hour since your claim time for '{item.name}'.",
            type="claim_late_smart",
            scheduled_at=reminder3_time,
        )

def pretty_priority(p: str) -> str:
    if not p:
        return "Low — General"
    return PRIORITY_LABELS.get(str(p), str(p))

   


def api_inventory_list(request):
    items = Item.objects.all().values(
        'item_id', 'name', 'qty', 'category', 'description', 'owner', 'status', 'image'
    )
    data = list(items)
    for item in data:
        if item['image']:
            item['image'] = request.build_absolute_uri(f"/media/{item['image']}")
        else:
            item['image'] = None
    return JsonResponse(data, safe=False, status=200)


def api_inventory_detail(request, id):
    try:
        item = Item.objects.get(item_id=id)
        data = {
            "item_id": item.item_id,
            "item_name": item.name,
            "description": item.description,
            "quantity": item.qty,
            "item_owner": item.owner or "Barangay Kauswagan",
            "item_image": request.build_absolute_uri(f"/media/{item.image}") if item.image else None,
        }
        return JsonResponse(data, status=200)
    except Item.DoesNotExist:
        raise Http404("Item not found")





def total_reserved_qty_for_range(item, start_date, end_date):
    """
    Calculate total quantity reserved for any overlapping date range.
    Includes reservations that overlap the target range and are pending/approved.
    """
    overlap_filter = Q(date_borrowed__lte=end_date, date_return__gte=start_date)
    agg = (
        Reservation.objects
        .filter(item=item, status__in=['pending', 'approved'])
        .filter(overlap_filter)
        .aggregate(total=Sum('quantity'))
    )
    return agg['total'] or 0


def find_next_available_dates(item, want_qty, start_date, horizon_days=30, limit=3):
    """Suggest next future ranges where the item can fit."""
    suggestions = []
    current = start_date
    while len(suggestions) < limit and current < start_date + timedelta(days=horizon_days):
        reserved = total_reserved_qty_for_range(item, current, current)
        total_stock = item.qty + (
            Reservation.objects.filter(item=item, status__in=['pending', 'approved'])
            .aggregate(total=Sum('quantity')).get('total', 0) or 0
        )
        if reserved + want_qty <= total_stock:
            suggestions.append({"date": current.isoformat()})
        current += timedelta(days=1)
    return suggestions



class CheckAvailabilityView(APIView):
    """
    Check if a single item has enough available quantity within a date range.
    Now fully compatible with multi-item reservations.
    """
    permission_classes = [AllowAny]

    def post(self, request):
        try:
            item_id = int(request.data.get("item_id"))
            want_qty = int(request.data.get("qty"))
            start_date = date.fromisoformat(request.data.get("start_date"))
            end_date = date.fromisoformat(request.data.get("end_date"))
        except Exception:
            return Response({"detail": "Invalid payload."}, status=400)

        if want_qty < 1:
            return Response({"detail": "Quantity must be >= 1"}, status=400)
        if start_date > end_date:
            return Response({"detail": "Invalid date range."}, status=400)

        # ITEM EXISTS?
        try:
            item = Item.objects.get(item_id=item_id)
        except Item.DoesNotExist:
            return Response({"detail": "Item not found."}, status=404)

        # ===== ADMIN BLOCK CHECK =====
        if BlockedDate.objects.filter(
            item=item,
            date__range=[start_date, end_date]
        ).exists():
            return Response({
                "detail": "This date range is blocked by the administrator.",
                "blocked": True,
            }, status=409)

        # ===== TOTAL RESERVED (using ReservationItem — NEW SYSTEM) =====
        overlap = ReservationItem.objects.filter(
            item=item,
            reservation__status__in=["approved", "in use"],
            reservation__date_borrowed__lte=end_date,
            reservation__date_return__gte=start_date,
        ).aggregate(total=Sum("quantity"))

        reserved = overlap["total"] or 0
        available = max(item.qty - reserved, 0)

        if available < want_qty:
            return Response({
                "detail": "Not enough items available for that range.",
                "available_qty": available,
            }, status=409)

        return Response({
            "ok": True,
            "available_qty": available
        }, status=200)

class CreateReservationView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser, JSONParser]

    @transaction.atomic
    def post(self, request):
        try:
            borrower = UserBorrower.objects.get(user=request.user)

            main_item_id = int(request.data.get("main_item_id"))
            main_item_qty = int(request.data.get("main_item_qty"))
            added_items = json.loads(request.data.get("added_items", "[]"))

            start_date = date.fromisoformat(request.data.get("start_date"))
            end_date = date.fromisoformat(request.data.get("end_date"))

            
            priority_detail = request.data.get("priority_detail", "")
            other_reason = request.data.get("other_reason", "")
            
            if priority_detail == "Others" and other_reason:
                priority_detail = other_reason

            # Define which reasons count as HIGH priority
            HIGH_CATEGORIES = ["Funeral", "Government Activities"]

            if priority_detail in HIGH_CATEGORIES:
                priority = "High"
            else:
                priority = "Low"
                
            message = request.data.get("message", "")

            letter_img = request.FILES.get("letter_image")
            valid_id_img = request.FILES.get("valid_id_image")

            contact_number = request.data.get("contact", borrower.contact_number)

        except Exception as e:
            return Response({"detail": "Invalid payload", "error": str(e)}, status=400)

        if start_date > end_date:
            return Response({"detail": "Invalid date range"}, status=400)

        # ---- CREATE RESERVATION ----
        reservation = Reservation.objects.create(
            userborrower=borrower,
            date_borrowed=start_date,
            date_return=end_date,
            priority=priority,
            priority_detail=priority_detail,
            message=message,
            status="pending",
            letter_image=letter_img,
            valid_id_image=valid_id_img,
            contact=contact_number,
        )

        reservation.save()

        # ------------------------------------------------------
        # INTERNAL VALIDATOR (DO NOT TOUCH)
        # ------------------------------------------------------
        def validate_and_add(item_id, qty):
            item = Item.objects.select_for_update().get(item_id=item_id)

            overlapping = ReservationItem.objects.filter(
                item=item,
                reservation__status__in=["approved", "in use"],
                reservation__date_borrowed__lte=end_date,
                reservation__date_return__gte=start_date,
            ).aggregate(total=Sum("quantity"))

            reserved_qty = overlapping["total"] or 0
            available = item.qty - reserved_qty

            if qty > available:
                raise ValueError(f"{item.name}: only {available} available.")

            ReservationItem.objects.create(
                reservation=reservation,
                item=item,
                item_name=item.name,
                quantity=qty
            )

        # ---- ADD MAIN ITEM ----
        try:
            validate_and_add(main_item_id, main_item_qty)
        except Exception as e:
            reservation.delete()
            return Response({"detail": str(e)}, status=409)

        # ---- ADD EXTRA ITEMS ----
        for it in added_items:
            try:
                validate_and_add(int(it["id"]), int(it["qty"]))
            except Exception as e:
                reservation.delete()
                return Response({"detail": str(e)}, status=409)

        # ------------------------------------------------------
        # ⭐ COMPUTE PRIORITY SCORE (FINAL FIX)
        # ------------------------------------------------------
        reservation.refresh_from_db()  # ensure items exist
        reservation.priority_score = compute_priority(reservation)
        reservation.save()

        # ------------------------------------------------------
        # SEND INITIAL NOTIFICATION (unchanged)
        # ------------------------------------------------------
        Notification.objects.create(
            user=borrower,
            reservation=reservation,
            title="Pending Reservation",
            message=f"Your reservation ({reservation.transaction_id}) is pending approval.",
            type="pending",
            is_sent=True
        )

        return Response({
            "success": True,
            "transaction_id": reservation.transaction_id,
            "priority_score": reservation.priority_score  # return priority to mobile
        }, status=201)

@csrf_exempt
def user_profile(request):
    """
    API endpoint to fetch user borrower profile
    """
    if request.method == "GET":
        try:
            username = request.GET.get("username")

            if not username:
                return JsonResponse({"success": False, "message": "Username is required"}, status=400)

            # Find the user
            user = User.objects.filter(username=username).first()
            if not user:
                return JsonResponse({"success": False, "message": "User not found"}, status=404)

            # Find the borrower profile
            borrower = UserBorrower.objects.filter(user=user).first()
            if not borrower:
                return JsonResponse({"success": False, "message": "No profile found"}, status=404)

            #  Get image URL (check if profile_image field exists)
            image_url = borrower.profile_image.url if getattr(borrower, "profile_image", None) else None

            # Return all user borrower data including image
            return JsonResponse({
                "success": True,
                "data": {
                    "username": user.username,
                    "name": borrower.full_name,
                    "contactNumber": borrower.contact_number,
                    "address": borrower.address,
                    "image": image_url,  # 👈 Added this line
                }
            }, status=200)

        except Exception as e:
            return JsonResponse({"success": False, "message": str(e)}, status=400)

    # Handle incorrect request methods
    return JsonResponse({"success": False, "message": "Invalid request method"}, status=405)




@csrf_exempt
def update_profile(request):
    if request.method == "POST":
        try:
            username = request.POST.get("username")
            name = request.POST.get("name")
            contact_number = request.POST.get("contactNumber")
            address = request.POST.get("address")
            password = request.POST.get("password")

            user = User.objects.filter(username=username).first()
            if not user:
                return JsonResponse({"success": False, "message": "User not found"}, status=404)

            borrower = UserBorrower.objects.filter(user=user).first()
            if not borrower:
                return JsonResponse({"success": False, "message": "Profile not found"}, status=404)

            # Update fields
            borrower.full_name = name
            borrower.contact_number = contact_number
            borrower.address = address

            # Handle image upload
            if "profile_image" in request.FILES:
                borrower.profile_image = request.FILES["profile_image"]

            borrower.save()

            if password:
                user.set_password(password)
                user.save()

            return JsonResponse({"success": True, "message": "Profile updated successfully"})
        except Exception as e:
            return JsonResponse({"success": False, "message": str(e)}, status=400)
    return JsonResponse({"success": False, "message": "Invalid request method"}, status=405)


# NEW

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def save_device_token(request):
    """Save or update the borrower's device token."""
    user = request.user
    token = request.data.get('token')

    if not token:
        return Response({'success': False, 'message': 'Token required'}, status=400)

    DeviceToken.objects.update_or_create(user=user, defaults={'token': token})
    return Response({'success': True, 'message': 'Token saved'})

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_user_notifications(request):
    """Return list of notifications for the logged-in borrower with multi-item support."""
    try:
        borrower = getattr(request.user, 'userborrower', None)
        if not borrower:
            return Response({'success': True, 'notifications': []}, status=200)

        # ⭐ FIX: Only return notifications that are ACTUALLY SENT
        notifications = (
            Notification.objects
            .filter(user=borrower, is_sent=True)
            .select_related('reservation')
            .order_by('-created_at')
        )


        ICONS = {
            "approval": "checkmark-circle-outline",
            "rejection": "close-circle-outline",
            "pending": "time-outline",
            "claimed": "cube-outline",
            "returned": "arrow-undo-outline",
            "cancelled": "close-circle-outline",
            "warning_claim_delay": "alert-outline",
            "claim_reminder": "time-outline",
            "return_reminder": "alert-circle-outline",
            "general": "notifications-outline"
        }

        data = []

        for n in notifications:
            reservation = n.reservation

            # ---------- MULTI-ITEM SUPPORT ----------
            item_list = []
            if reservation:
                for it in reservation.items.all():
                    item_list.append({
                        "item_id": it.item.item_id,
                        "item_name": it.item_name,
                        "quantity": it.quantity,
                        "image": (
                            request.build_absolute_uri(it.item.image.url)
                            if it.item.image else None
                        )
                    })

            # Convert to local timezone
            local_time = timezone.localtime(n.created_at)

            data.append({
                "id": n.id,
                "title": n.title,
                "message": n.message,
                "reason": n.reason,
                "type": n.type,
                "icon": ICONS.get(n.type, "notifications-outline"),
                "is_read": n.is_read,
                "created_at": local_time.strftime("%Y-%m-%d %I:%M %p"),

                # QR
                "qr_code": (
                    request.build_absolute_uri(n.qr_code.url)
                    if n.qr_code else None
                ),

                # MULTI-ITEM + TRANSACTION
                "transaction_id": reservation.transaction_id if reservation else None,
                "items": item_list,
            })

        return Response({"success": True, "notifications": data}, status=200)

    except Exception as e:
        import traceback; traceback.print_exc()
        return Response({"success": False, "error": str(e)}, status=500)

    

def create_notification(borrower, title, message, notif_type='general', qr_file=None, reservation=None):
    """Reusable helper to create both in-app + push notification."""
    notif = Notification.objects.create(
        user=borrower,
        reservation=reservation, 
        title=title,
        message=message,
        type=notif_type
    )

    if qr_file:
        notif.qr_code.save(f"qr_{borrower.user.username}.png", qr_file)

    # Optional push notification
    try:
        token_entry = DeviceToken.objects.filter(user=borrower.user).last()
        if token_entry:
            push_message = messaging.Message(
                notification=messaging.Notification(
                    title=title,
                    body=message
                ),
                token=token_entry.token,
            )
            messaging.send(push_message)
    except Exception as e:
        print("Error sending push notification:", e)

    return notif

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def add_delayed_notification(request):
    """
    Create a new 'delayed' notification for a borrower.
    Expected JSON:
    {
        "user_id": 5,
        "item_name": "Projector",
        "message": "You returned the item late. Please be punctual next time."
    }
    """
    try:
        user_id = request.data.get("user_id")
        message = request.data.get("message", "")
        item_name = request.data.get("item_name", "")

        if not user_id:
            return Response({"success": False, "message": "Missing user_id"}, status=400)

        borrower = UserBorrower.objects.get(id=user_id)

        title = "Delayed Return Notice"
        full_message = message or f"You returned '{item_name}' late. Please avoid future delays."

        # Create the delayed notification
        notif = Notification.objects.create(
            user=borrower,
            title=title,
            message=full_message,
            type="delayed",
        )

        # Optional: push notification
        try:
            token_entry = DeviceToken.objects.filter(user=borrower.user).last()
            if token_entry:
                push_message = messaging.Message(
                    notification=messaging.Notification(
                        title=title,
                        body=full_message[:200],
                    ),
                    token=token_entry.token,
                )
                messaging.send(push_message)
        except Exception as e:
            print("Push notification failed:", e)

        return Response({
            "success": True,
            "message": "Delayed notification sent successfully",
            "notification_id": notif.id
        }, status=201)

    except UserBorrower.DoesNotExist:
        return Response({"success": False, "message": "User not found"}, status=404)
    except Exception as e:
        import traceback; traceback.print_exc()
        return Response({"success": False, "message": str(e)}, status=500)



@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_notification(request, pk):
    """
    Permanently deletes a single notification.
    Only the owner of the notification can delete it.
    """
    try:
        notif = Notification.objects.get(pk=pk, user__user=request.user)
        notif.delete()
        return Response({"success": True, "message": "Notification permanently deleted"}, status=200)
    except Notification.DoesNotExist:
        return Response({"success": False, "message": "Notification not found"}, status=404)



@api_view(['PATCH'])
@permission_classes([IsAuthenticated])
def mark_notification_as_read(request, pk):
    """
    Marks a specific notification as read (is_read=True)
    """
    try:
        notif = Notification.objects.get(pk=pk, user__user=request.user)
        notif.is_read = True
        notif.save()
        return Response({'success': True, 'message': 'Notification marked as read'})
    except Notification.DoesNotExist:
        return Response({'success': False, 'message': 'Notification not found'}, status=404)

@api_view(['PATCH'])
@permission_classes([IsAuthenticated])
def mark_all_notifications_as_read(request):
    queryset = Notification.objects.filter(user__user=request.user, is_read=False)
    count = queryset.update(is_read=True)
    return Response({
        'success': True,
        'message': f'{count} notifications marked as read'
    })   

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def trigger_due_soon_notifications(request):
    send_due_soon_notifications()
    return Response({'success': True, 'message': 'Return reminders sent successfully'})

def send_due_soon_notifications():
    """
    Creates 'Return Reminder' notifications for items due tomorrow.
    Can be called manually or via a scheduled task.
    """
    today = date.today()
    tomorrow = today + timedelta(days=1)

    due_soon = Reservation.objects.filter(
        date_return=tomorrow,
        status__in=['approved', 'in use']
    ).select_related('userborrower', 'item')

    for r in due_soon:
        borrower = r.userborrower
        item = r.item

        if not borrower or not item:
            continue

        # Avoid duplicates
        already_sent = Notification.objects.filter(
            user=borrower,
            reservation=r,
            type='return_reminder'
        ).exists()
        if already_sent:
            continue

        Notification.objects.create(
            user=borrower,
            reservation=r,
            title="Return Reminder",
            message=f"Your borrowed item '{item.name}' is due for return tomorrow. Please return it on time to avoid penalties.",
            type="return_reminder",
        )
        
        
        
        
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def user_reservations(request):
    borrower = getattr(request.user, 'userborrower', None)
    if not borrower:
        return Response({'success': False, 'reservations': []})

    reservations = Reservation.objects.filter(userborrower=borrower).order_by('-created_at')

    data = []
    for r in reservations:

        items_list = []
        reserved_items = r.items.all()   # <-- correct related_name

        for ri in reserved_items:
            items_list.append({
                "name": ri.item_name,
                "quantity": ri.quantity,
                "image_url": (
                    request.build_absolute_uri(ri.item.image.url)
                    if ri.item and ri.item.image else None
                )
            })

        data.append({
            'id': r.id,
            'transaction_id': r.transaction_id,
            'status': r.status,
            'priority': r.priority,
            'date_borrowed': r.date_borrowed.strftime('%Y-%m-%d') if r.date_borrowed else None,
            'date_return': r.date_return.strftime('%Y-%m-%d') if r.date_return else None,
            "reason_for_borrowing": r.priority_detail,
            'items': items_list,
        })

    return Response({'success': True, 'reservations': data}, status=200)


        

@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def cancel_reservation(request, pk):
    """
    Allows the borrower to cancel their own pending reservation.
    """

    borrower = getattr(request.user, 'userborrower', None)
    if not borrower:
        return Response({'success': False, 'message': 'Unauthorized user.'}, status=403)

    try:
        reservation = Reservation.objects.get(pk=pk, userborrower=borrower)

        if reservation.status != 'pending':
            return Response({'success': False, 'message': 'Only pending reservations can be cancelled.'}, status=400)

        # NEW — Extract correct item names from ReservationItem
        item_names = [ri.item_name for ri in reservation.items.all()]
        item_display = ", ".join(item_names) if item_names else "your items"

        # Update status
        reservation.status = 'cancelled'
        reservation.save()

        # Create Cancellation Notification
        create_notification(
            borrower,
            title="Reservation Cancelled",
            message=f"You cancelled your reservation for {item_display}.",
            notif_type="cancelled"
        )

        return Response({'success': True, 'message': 'Reservation cancelled successfully.'}, status=200)

    except Reservation.DoesNotExist:
        return Response({'success': False, 'message': 'Reservation not found.'}, status=404)

# NEW — Dynamic availability for a single date
# ----------------------------------------------------------
# ITEM AVAILABILITY (Single Date)
# ----------------------------------------------------------
@api_view(["GET"])
@permission_classes([AllowAny])
def item_availability(request, item_id):
    from datetime import datetime, timedelta

    date_str = request.GET.get("date")
    if not date_str:
        return Response({"error": "Missing date parameter"}, status=400)

    try:
        selected_date = datetime.strptime(date_str, "%Y-%m-%d").date()
    except ValueError:
        return Response({"error": "Invalid date format"}, status=400)

    try:
        item = Item.objects.get(pk=item_id)
    except Item.DoesNotExist:
        return Response({"error": "Item not found"}, status=404)

    total_capacity = item.qty

    overlapping = ReservationItem.objects.filter(
        item=item,
        reservation__status__in=["approved", "in use"],
        reservation__date_borrowed__lte=selected_date,
        reservation__date_return__gte=selected_date,
    ).aggregate(total=Sum("quantity"))

    reserved = overlapping["total"] or 0
    available_qty = max(total_capacity - reserved, 0)
    status = "fully_reserved" if available_qty == 0 else "available"

    # Suggested next date
    suggested_date = None
    if status == "fully_reserved":
        next_day = selected_date + timedelta(days=1)
        for _ in range(30):
            overlap_next = ReservationItem.objects.filter(
                item=item,
                reservation__status__in=["pending", "approved", "in use"],
                reservation__date_borrowed__lte=next_day,
                reservation__date_return__gte=next_day,
            ).aggregate(total=Sum("quantity"))["total"] or 0

            if total_capacity - overlap_next > 0:
                suggested_date = next_day.isoformat()
                break
            next_day += timedelta(days=1)

    return Response({
        "item_id": item.item_id,
        "item_name": item.name,
        "date": selected_date.isoformat(),
        "status": status,
        "available_qty": available_qty,
        "suggested_date": suggested_date,
    })



# ----------------------------------------------------------
# ITEM AVAILABILITY MAP (60-day range)
# ----------------------------------------------------------
@api_view(["GET"])
@permission_classes([AllowAny])
def item_availability_map(request, item_id):
    from datetime import date, timedelta

    try:
        item = Item.objects.get(pk=item_id)
    except Item.DoesNotExist:
        return Response({"error": "Item not found"}, status=404)

    total_capacity = item.qty

    res_items = ReservationItem.objects.filter(
        item=item,
        reservation__status__in=["approved", "in use"]
    ).select_related("reservation")

    reservations = [
        {
            "start": ri.reservation.date_borrowed,
            "end": ri.reservation.date_return,
            "quantity": ri.quantity,
            "borrower": ri.reservation.userborrower.full_name if ri.reservation.userborrower else "Unknown",
            "status": ri.reservation.status,
        }
        for ri in res_items
    ]

    admin_borrows = list(
        AdminBorrow.objects.filter(
            item=item,
            status="In Use"
        ).values("date", "return_date", "quantity", "borrower_name",
                 "contact_number", "address", "purpose", "delivered_by", "id", "transaction_id")
    )

    blocked_dates = set(b.date for b in BlockedDate.objects.filter(item=item))

    start = date.today()
    end = start + timedelta(days=60)
    days = {}
    current = start

    while current <= end:
        key = current.isoformat()

        if current in blocked_dates:
            days[key] = {
                "status": "blocked",
                "reserved_qty": 0,
                "admin_borrowed": 0,
                "available_qty": 0,
            }
            current += timedelta(days=1)
            continue

        reserved_qty = 0
        daily_reservations = []
        for r in reservations:
            if r["start"] <= current <= r["end"]:
                reserved_qty += r["quantity"]
                daily_reservations.append(r)

        admin_used = 0
        admin_full = []
        for ab in admin_borrows:
            if ab["date"] <= current <= ab["return_date"]:
                admin_used += ab["quantity"]
                admin_full.append(ab)

        available = max(total_capacity - reserved_qty - admin_used, 0)

        days[key] = {
            "status": "fully_reserved" if available == 0 else "available",
            "reserved_qty": reserved_qty,
            "admin_borrowed": admin_used,
            "available_qty": available,
            "reservations": daily_reservations,
            "admin_borrow_full": admin_full,
        }

        current += timedelta(days=1)

    return Response({
        "item_id": item.item_id,
        "item_name": item.name,
        "calendar": days
    })




def get_total_capacity(item):
    return item.qty or 0


@csrf_exempt
def verify_qr(request, mode, code):
    try:
        # Extract T + 6 digits from FULL QR TEXT
        match = re.search(r"T\d{6}", code)
        if not match:
            return JsonResponse({"error": "Invalid QR format"}, status=400)

        code = match.group(0)

        r = Reservation.objects.prefetch_related("items__item").get(transaction_id=code)

        # FRONTEND-FRIENDLY STATUS
        status_map = {
            "approved": "Approved",
            "in use": "In Use",
            "returned": "Returned",
            "pending": "Pending",
            "declined": "Declined",
            "cancelled": "Cancelled"
        }
        ui_status = status_map.get(r.status.lower(), r.status)

        data = {
            "borrower": r.userborrower.full_name,
            "status": ui_status,
            "start": r.date_borrowed,
            "end": r.date_return,
            "delivered_by": r.delivered_by,

            # Return UI needs these:
            "can_claim": (mode == "claim" and r.status.lower() == "approved"),
            "can_return": (mode == "return" and r.status.lower() == "in use"),

            "items": [
                {"name": it.item_name, "qty": it.quantity}
                for it in r.items.all()
            ]
        }

        return JsonResponse(data)

    except Reservation.DoesNotExist:
        return JsonResponse({"error": "QR not recognized"}, status=404)

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)


    
@csrf_exempt
def update_reservation(request, mode, code):
    try:
        # Extract transaction ID
        match = re.search(r"T\d{6}", code)
        if not match:
            return JsonResponse({"error": "Invalid QR format"}, status=400)

        code = match.group(0)

        reservation = Reservation.objects.prefetch_related("items").get(transaction_id=code)
        delivered_by = request.GET.get("delivered_by", "").strip()

        # -------------------------
        # CLAIM MODE
        # -------------------------
        if mode.lower() == "claim":

            if reservation.status.lower() == "in use":
                return JsonResponse({"error": "This reservation is already claimed."}, status=400)

            if reservation.status.lower() != "approved":
                return JsonResponse({"error": "This reservation is not approved for claiming."}, status=400)

            if not delivered_by:
                return JsonResponse({"error": "Delivered By is required."}, status=400)

            # ==============================
            # FIXED: CLAIM DATE UPDATES
            # ==============================
            today = timezone.now().date()
            reservation.status = "in use"
            reservation.date_receive = timezone.now()

            reservation.borrow_date = today
            reservation.actual_claimed_date = today

            # Fill expected return date if empty
            reservation.expected_return_date = reservation.date_return

            reservation.delivered_by = delivered_by
            reservation.save()

            item_list = ", ".join([i.item_name for i in reservation.items.all()])

            # Notify
            Notification.objects.create(
                user=reservation.userborrower,
                reservation=reservation,
                title="Item Claimed",
                message=f"You have claimed: {item_list}",
                type="claimed",
                is_sent=True
            )

            return JsonResponse({"message": "Claim successful"})

        # -------------------------
        # RETURN MODE
        # -------------------------
        elif mode.lower() == "return":

            if reservation.status.lower() != "in use":
                return JsonResponse({"error": "This item cannot be returned."}, status=400)

            # Do not finalize here — submit_feedback does that
            return JsonResponse({"message": "Return validated. Proceed to form."})

        else:
            return JsonResponse({"error": "Invalid mode"}, status=400)

    except Reservation.DoesNotExist:
        return JsonResponse({"error": "Reservation not found"}, status=404)

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)



@csrf_exempt
def submit_feedback(request):
    try:
        import re

        if request.method != "POST":
            return JsonResponse({"error": "Invalid request"}, status=400)

        # ------------------------------------------
        # CLEAN TRANSACTION ID
        # ------------------------------------------
        raw_id = request.POST.get("transaction_id", "")
        match = re.search(r"T\d{6}", raw_id)
        if not match:
            return JsonResponse({"error": "Invalid transaction ID"}, status=400)
        transaction_id = match.group(0)

        # BASIC FIELDS
        comment = request.POST.get("comment", "").strip()
        return_status = request.POST.get("return_status", "").strip()  # None / Damage / Loss / On Time / Late

        # ISSUE-TYPE FIELDS
        damage_qty = request.POST.get("damage_qty")
        damage_area = request.POST.get("damage_area")
        damage_description = request.POST.get("damage_description")

        loss_qty = request.POST.get("loss_qty")
        loss_area = request.POST.get("loss_area")
        loss_reason = request.POST.get("loss_reason")

        return_qty = request.POST.get("return_qty", "0")

        # ------------------------------------------
        # FETCH RESERVATION + BORROWER
        # ------------------------------------------
        reservation = Reservation.objects.prefetch_related("items__item").get(transaction_id=transaction_id)
        borrower = reservation.userborrower
        items = reservation.items.all()
        item = items.first().item

        total_qty = sum(i.quantity for i in items)

        # ------------------------------------------
        # VALIDATE RETURN QUANTITY
        # ------------------------------------------
        try:
            return_qty = int(return_qty)
            if return_qty < 0 or return_qty > total_qty:
                return JsonResponse({"error": "Invalid return quantity"}, status=400)
        except:
            return JsonResponse({"error": "Return quantity must be numeric"}, status=400)

        missing_qty = total_qty - return_qty

        # ------------------------------------------
        # CREATE FEEDBACK ENTRY
        # ------------------------------------------
        fb = Feedback.objects.create(
            reservation=reservation,
            userborrower=borrower,
            comment=comment,
            return_status=return_status
        )

        # ------------------------------------------
        # AUTO COMPUTE RETURN TIMELINE
        # ------------------------------------------
        today = timezone.now().date()
        due_date = reservation.date_return

        if today < due_date:
            auto_status = "Early Return"
            late_days = 0
        elif today == due_date:
            auto_status = "On Time"
            late_days = 0
        else:
            late_days = (today - due_date).days
            auto_status = f"Late Return ({late_days} days)"

        fb.days_late = late_days
        fb.loss_qty = missing_qty
        fb.save()

        # ------------------------------------------
        # VIOLATION SCORING
        # ------------------------------------------
        violation_points = 0

        if late_days > 0:
            if late_days <= 3:
                violation_points += 1
            elif late_days <= 7:
                violation_points += 3
            else:
                violation_points += 5

        # MISSING ITEMS
        if return_status == "Missing" and missing_qty > 0:
            violation_points += 5

        # ------------------------------------------
        # DAMAGE CASE
        # ------------------------------------------
        if return_status == "Damage":
            if not damage_qty or not damage_area or not damage_description:
                return JsonResponse({"error": "Damage details incomplete"}, status=400)

            DamageReport.objects.create(
                reported_by=borrower,
                reservation=reservation,
                item=item,
                report_type="Damage",
                quantity_affected=int(damage_qty),
                location=damage_area.strip(),
                description=damage_description.strip()
            )

            reservation.report_type = "Damage"
            violation_points += 4

        # ------------------------------------------
        # LOSS CASE
        # ------------------------------------------
        elif return_status == "Loss":
            if not loss_qty or not loss_reason:
                return JsonResponse({"error": "Loss details incomplete"}, status=400)

            DamageReport.objects.create(
                reported_by=borrower,
                reservation=reservation,
                item=item,
                report_type="Loss",
                quantity_affected=int(loss_qty),
                location=loss_area if loss_area else "N/A",
                description=loss_reason.strip()
            )

            item.qty = max(0, item.qty - int(loss_qty))
            item.save()

            reservation.report_type = "Loss"
            violation_points += 6

        else:
            # NONE
            reservation.report_type = "None"

        # ------------------------------------------
        # UPDATE BORROWER SCORE
        # ------------------------------------------
        borrower.behavior_total_score += violation_points
        borrower.behavior_score = borrower.behavior_total_score
        borrower.violation_count += 1

        if borrower.behavior_total_score < 5:
            borrower.borrower_status = "Good"
            notif_type = "returned"
            notif_title = "Return Successful"
            notif_message = "Thank you for returning your borrowed items."

        elif borrower.behavior_total_score < 10:
            borrower.borrower_status = "Warning"
            notif_type = "warning"
            notif_title = "Warning Issued"
            notif_message = "You now have multiple violations."

        else:
            borrower.borrower_status = "Bad"
            notif_type = "restricted"
            notif_title = "Account Restricted"
            notif_message = "Your borrowing privileges are now restricted."

        borrower.save()

        # ------------------------------------------
        # FINALIZE RESERVATION TIMELINE
        # ------------------------------------------
        if not reservation.expected_return_date:
            reservation.expected_return_date = reservation.date_return

        reservation.actual_return_date = timezone.now().date()
        reservation.date_returned = timezone.now()
        reservation.missing_qty = missing_qty
        reservation.returned_qty = return_qty

        reservation.status = "returned"
        reservation.save()

        # ------------------------------------------
        # SEND NOTIFICATION
        # ------------------------------------------
        Notification.objects.create(
            user=borrower,
            reservation=reservation,
            title=notif_title,
            message=notif_message,
            type=notif_type,
            is_sent=True
        )

        return JsonResponse({
            "message": "Return processed successfully",
            "late_days": late_days,
            "missing_qty": missing_qty,
            "behavior_score": borrower.behavior_total_score
        })

    except Reservation.DoesNotExist:
        return JsonResponse({"error": "Reservation not found"}, status=404)

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)

        # ============================================================
# BEHAVIOR PATTERN DETECTION ENGINE
# ============================================================
def analyze_borrower_patterns(borrower):
    """
    Auto-adjusts behavior_score based on long-term patterns.
    This runs after every feedback submission.
    """

    # ---------- 1. Count last 5 returns ----------
    recent_feedback = Feedback.objects.filter(
        userborrower=borrower
    ).order_by('-created_at')[:5]

    late_count = sum(1 for fb in recent_feedback if fb.return_status == "Late Return")
    on_time_count = sum(1 for fb in recent_feedback if fb.return_status == "On Time")
    not_returned_count = sum(1 for fb in recent_feedback if fb.return_status == "Not Returned")

    # ---------- 2. Count damage/loss reports ----------
    damage_reports = DamageReport.objects.filter(
        reported_by=borrower
    ).count()

    loss_reports = DamageReport.objects.filter(
        reported_by=borrower,
        report_type="Loss"
    ).count()

    behavior_change = 0

    # =====================================================
    # PATTERN 1: Frequent Late Returns
    # =====================================================
    if late_count >= 3:
        # borrower is repeatedly late
        behavior_change -= 2  

    # =====================================================
    # PATTERN 2: Multiple Damages or Losses
    # =====================================================
    if damage_reports >= 2:
        behavior_change -= 3  # repeated damages

    if loss_reports >= 1:
        behavior_change -= 4  # loss is a major violation

    # =====================================================
    # PATTERN 3: Perfect On-Time streak (reward)
    # =====================================================
    if on_time_count >= 5:
        behavior_change += 3  # encourage good borrowers

    # =====================================================
    # APPLY BEHAVIOR CHANGES
    # =====================================================
    if behavior_change != 0:
        borrower.behavior_score += behavior_change

        # Update borrower status based on thresholds
        if borrower.behavior_score >= 0:
            borrower.borrower_status = "Good"

        elif -9 <= borrower.behavior_score < 0:
            borrower.borrower_status = "Warning"

        else:
            borrower.borrower_status = "Bad"

        borrower.save()



    return behavior_change




@csrf_exempt
def monthly_reset(request=None):
    """
    Smart monthly reset:
    - Clears late_count
    - Gives small behavior recovery (if not Bad)
    - Resets Warning borrowers back to Good
    - Bad borrowers remain restricted
    """
    try:
        borrowers = UserBorrower.objects.all()

        for b in borrowers:

            # 1️⃣ Reset late count every month
            b.late_count = 0

            # 2️⃣ Slight positive recovery for Good/Warning users
            if b.borrower_status in ["Good", "Warning"]:
                b.behavior_score += 2

            # 3️⃣ Warning -> Good
            if b.borrower_status == "Warning":
                b.borrower_status = "Good"

            # 4️⃣ If behavior_score recovered above 0 → Good
            if b.behavior_score >= 0 and b.borrower_status != "Bad":
                b.borrower_status = "Good"

            # Bad stays Bad — must be manually reviewed
            b.save()

        return JsonResponse({"message": "Monthly reset completed successfully."})

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)


def damage_loss_report_list(request):
    reports = DamageReport.objects.select_related('reported_by').order_by('-date_reported')

    report_data = []
    for r in reports:
        local_time = timezone.localtime(r.date_reported)

        report_data.append({
            'id': r.id,
            'user_id': r.reported_by.id,
            'user_name': r.reported_by.full_name,
            'address': r.reported_by.address,
            'type': r.report_type,   
            'image': r.image.url if r.image else 'No image',
            'date': local_time.strftime("%Y-%m-%d %I:%M %p"),
            'description': r.description,
            'quantity': r.quantity_affected,
            'location': r.location,
            'status': r.status,
            'item_name': r.item.name if r.item else "Unknown Item"
        })

    return render(request, 'damage_report.html', {'reports': report_data})



@api_view(['POST'])
@permission_classes([IsAuthenticated])
@parser_classes([MultiPartParser, FormParser])
def submit_damage_loss_report(request):
    """
    Borrower submits a Damage or Loss report for an item currently in use.
    Links to reservation + deducts item qty for Loss reports.
    """
    try:
        borrower = UserBorrower.objects.get(user=request.user)

        transaction_id = request.data.get("reservation_id")
        item_id = request.data.get("item_id")
        report_type = request.data.get("report_type")
        location = request.data.get("location")
        quantity_affected = int(request.data.get("quantity_affected"))
        description = request.data.get("description")
        image = request.data.get("image")

        # Validate
        if not all([transaction_id, item_id, report_type, location, description]):
            return Response({"success": False, "message": "Missing required fields"}, status=400)

        reservation = Reservation.objects.get(id=transaction_id, userborrower=borrower)
        item = Item.objects.get(item_id=item_id)

        # Create the report
        report = DamageReport.objects.create(
            reported_by=borrower,
            report_type=report_type,
            location=location,
            quantity_affected=quantity_affected,
            description=description,
            image=image,
            item=item,                
            reservation=reservation,  
        )


        # -------------------------
        # AUTO DEDUCT FOR LOSS ONLY
        # -------------------------
        if report_type.lower() == "loss":
            if item.qty >= quantity_affected:
                item.qty -= quantity_affected
            else:
                item.qty = 0     # never negative
            item.save()

        return Response({
            "success": True,
            "message": f"{report_type} report submitted successfully!",
            "report_id": report.id
        })

    except Reservation.DoesNotExist:
        return Response({"success": False, "message": "Reservation not found"}, status=404)

    except Item.DoesNotExist:
        return Response({"success": False, "message": "Item not found"}, status=404)

    except Exception as e:
        return Response({"success": False, "message": str(e)}, status=500)


@api_view(['GET'])
@authentication_classes([JWTAuthentication])
@permission_classes([IsAuthenticated])
def get_in_use_items(request):
    borrower = UserBorrower.objects.filter(user=request.user).first()
    if not borrower:
        return Response({"success": True, "items": []}, status=200)

    reservations = (
        Reservation.objects
        .filter(userborrower=borrower, status__iexact='in use')
        .prefetch_related('items__item')
        .order_by('-date_borrowed')
    )

    data = []

    for r in reservations:
        data.append({
            "reservation_id": r.id,
            "transaction_id": r.transaction_id,
            "date_borrowed": r.date_borrowed.strftime("%Y-%m-%d"),
            "date_return": r.date_return.strftime("%Y-%m-%d") if r.date_return else None,
            "items": [
                {
                    "item_id": ri.item.item_id,
                    "item_name": ri.item.name,
                    "quantity": ri.quantity,
                    "image": request.build_absolute_uri(ri.item.image.url) if ri.item.image else None
                }
                for ri in r.items.all()
            ]
        })

    return Response({"success": True, "items": data})



# ITEM CALENDAR BLOCKDATE

# ----------------------------------------------------------
# FULLCALENDAR WEB COMPATIBLE CALENDAR
# ----------------------------------------------------------
@api_view(["GET"])
@permission_classes([AllowAny])
def get_item_calendar(request, item_id):
    from datetime import timedelta, date

    try:
        item = Item.objects.get(item_id=item_id)
    except Item.DoesNotExist:
        return Response({"error": "Item not found"}, status=404)

    total_capacity = item.qty

    res_items = ReservationItem.objects.filter(
        item=item,
        reservation__status__in=["pending", "approved", "in use"]
    ).select_related("reservation", "reservation__userborrower")

    flat_reservations = []
    for ri in res_items:
        r = ri.reservation
        flat_reservations.append({
            "name": r.userborrower.full_name if r.userborrower else "Unknown",
            "date_borrowed": r.date_borrowed.isoformat(),
            "date_return": r.date_return.isoformat(),
            "quantity": ri.quantity,
            "status": r.status,
        })

    admin_borrow_qs = AdminBorrow.objects.filter(item=item, status="In Use")
    flat_admin = []
    for ab in admin_borrow_qs:
        flat_admin.append({
            "start": ab.date.isoformat(),
            "end": ab.return_date.isoformat(),
            "quantity": ab.quantity,
            "borrower_name": ab.borrower_name,
            "contact_number": ab.contact_number,
            "address": ab.address,
            "purpose": ab.purpose,
            "delivered_by": ab.delivered_by,
            "transaction_id": ab.transaction_id,
            "id": ab.id,
        })

    blocked = list(BlockedDate.objects.filter(item=item).values_list("date", flat=True))
    blocked_set = set(blocked)

    today = date.today()
    end = today + timedelta(days=60)

    reservations_by_date = {}
    admin_borrow_full = {}

    current = today
    while current <= end:
        key = current.isoformat()

        res_for_day = [r for r in flat_reservations if r["date_borrowed"] <= key <= r["date_return"]]
        reservations_by_date[key] = res_for_day

        admin_for_day = [ab for ab in flat_admin if ab["start"] <= key <= ab["end"]]
        admin_borrow_full[key] = admin_for_day

        current += timedelta(days=1)

    return Response({
        "item_id": item.item_id,
        "item_name": item.name,
        "reservations": [
            {"date": r["date_borrowed"], "name": r["name"], "status": r["status"]}
            for r in flat_reservations
        ],
        "admin_borrow": admin_borrow_full,
        "blocked": [d.isoformat() for d in blocked],
        "reservations_by_date": reservations_by_date,
        "admin_borrow_full": admin_borrow_full,
    })





# Unified block/unblock (for admin dashboard)
# ----------------------------------------------------------
# BLOCK / UNBLOCK a date (admin use)
# ----------------------------------------------------------
@csrf_exempt
@api_view(["POST"])
@permission_classes([AllowAny])
def toggle_block_date(request, item_id):
    try:
        body = json.loads(request.body or "{}")
        date_str = body.get("date")
        reason = body.get("reason", "Blocked manually")

        if not date_str:
            return Response({"error": "Missing date"}, status=400)

        date = parse_date(date_str)
        if not date:
            return Response({"error": "Invalid date format"}, status=400)

        item = Item.objects.get(item_id=item_id)

        existing = BlockedDate.objects.filter(item=item, date=date).first()
        if existing:
            existing.delete()
            return Response({"status": "unblocked", "date": date_str})

        BlockedDate.objects.create(item=item, date=date, reason=reason)
        return Response({"status": "blocked", "date": date_str, "reason": reason})

    except Item.DoesNotExist:
        return Response({"error": "Item not found"}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=500)

    
    


#  Cancel Reservation, qty increase and Notification
# ----------------------------------------------------------
# CANCEL ALL RESERVATIONS FOR A GIVEN DATE (ADMIN)
# ----------------------------------------------------------
@api_view(["POST"])
def cancel_reservations_for_date(request, item_id):
    try:
        date_str = request.data.get("date")
        if not date_str:
            return Response({"error": "Missing date"}, status=400)

        date = parse_date(date_str)
        if not date:
            return Response({"error": "Invalid date format"}, status=400)

        item = Item.objects.get(item_id=item_id)

        reservations = Reservation.objects.filter(
            item=item,
            date_borrowed__lte=date,
            date_return__gte=date
        ).exclude(status="cancelled")

        if not reservations.exists():
            return Response({"message": "No active reservations found for this date."}, status=200)

        total_restored = 0
        cancelled_count = 0

        for r in reservations:
            borrower = r.userborrower
            total_restored += r.quantity or 0
            r.status = "cancelled"
            r.save(update_fields=["status"])
            cancelled_count += 1

            if borrower:
                create_notification(
                    borrower,
                    title="Reservation Cancelled by Admin",
                    message=f"Your reservation for {r.item.name} has been cancelled.",
                    notif_type="cancelled",
                    reservation=r
                )

        item.qty += total_restored
        item.save(update_fields=["qty"])

        return Response({
            "message": f"{cancelled_count} reservation(s) cancelled.",
            "restored_qty": total_restored,
            "new_item_qty": item.qty
        }, status=200)

    except Item.DoesNotExist:
        return Response({"error": "Item not found"}, status=404)
    except Exception as e:
        return Response({"error": str(e)}, status=500)


    
    
@login_required
def change_password(request):
    if request.method == 'POST':
        current_password = request.POST.get('current_password')
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')

        user = request.user

        # Validate current password
        if not check_password(current_password, user.password):
            messages.error(request, "Current password is incorrect.")
            return redirect('change_password')

        # Validate new passwords match
        if new_password != confirm_password:
            messages.error(request, "New passwords do not match.")
            return redirect('change_password')

        # Validate password length (optional)
        if len(new_password) < 8:
            messages.error(request, "New password must be at least 8 characters long.")
            return redirect('change_password')

        # Save new password
        user.set_password(new_password)
        user.save()

        # Keep user logged in after password change
        update_session_auth_hash(request, user)

        messages.success(request, "Password updated successfully!")
        return redirect('change_password')

    return render(request, 'change_password.html')

#NEW
@csrf_exempt
def forgot_password(request):
    show_code_container = False
    email_value = ""  # Store email to keep it in the input

    if request.method == 'POST':
        # === SEND RESET CODE ===
        if 'send_code' in request.POST:
            email = request.POST.get('email')
            email_value = email

            try:
                user = User.objects.get(email=email)
                code = random.randint(100000, 999999)
                request.session['reset_email'] = email
                request.session['reset_code'] = str(code)

                # Build formal HTML email
                subject = "🔐 TrailLend Password Reset Code"
                from_email = settings.DEFAULT_FROM_EMAIL
                to = [email]

                html_content = f"""
                <!DOCTYPE html>
                <html lang="en">
                <head>
                  <meta charset="UTF-8" />
                  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
                  <title>Password Reset - TrailLend</title>
                </head>
                <body style="margin:0; padding:0; background-color:#f4f6f9; font-family:'Poppins',Arial,sans-serif;">
                  <table align="center" border="0" cellpadding="0" cellspacing="0" width="100%" 
                         style="max-width:600px; background-color:#ffffff; border-radius:10px; overflow:hidden; 
                                box-shadow:0 4px 10px rgba(0,0,0,0.05); margin-top:40px;">
                    <tr>
                      <td style="background-color:#1976D2; text-align:center; padding:30px;">
                        <img src="https://i.ibb.co/T2Hyfdd/TRAILLEND-ICON.png" alt="TrailLend Logo" width="80" style="margin-bottom:10px;" />
                        <h1 style="color:#fff; font-size:22px; margin:0;">TrailLend</h1>
                        <p style="color:#cce6ff; font-size:13px; margin:4px 0 0;">Empowering the Community Together 🌿</p>
                      </td>
                    </tr>
                    <tr>
                      <td style="padding:30px;">
                        <h2 style="color:#1976D2; font-size:20px; margin-top:0;">Password Reset Request</h2>
                        <p style="color:#333; font-size:15px;">Hello,</p>
                        <p style="color:#333; font-size:15px;">We received a request to reset your TrailLend account password. 
                        Please use the verification code below to continue:</p>

                        <div style="text-align:center; margin:30px 0;">
                          <span style="display:inline-block; background:#1976D2; color:#fff; font-size:28px; 
                                       letter-spacing:5px; padding:15px 25px; border-radius:8px; font-weight:bold;">
                            {code}
                          </span>
                        </div>

                        <p style="color:#555; font-size:14px;">Enter this code in the TrailLend app to verify your identity. 
                        This code will expire soon for security reasons.</p>

                        <p style="color:#555; font-size:14px;">If you did not request a password reset, please ignore this email. 
                        Your account remains secure.</p>

                        <hr style="border:none; border-top:1px solid #eee; margin:30px 0;" />

                        <p style="color:#999; font-size:12px; text-align:center;">
                          This email was sent by <strong>TrailLend</strong> • Barangay General Services Office<br/>
                          Please do not reply directly to this message.
                        </p>
                      </td>
                    </tr>
                  </table>
                </body>
                </html>
                """

                text_content = strip_tags(html_content)

                email_message = EmailMultiAlternatives(subject, text_content, from_email, to)
                email_message.attach_alternative(html_content, "text/html")
                email_message.send()

                messages.success(request, "A reset code has been sent to your email. Please check your inbox.")
                show_code_container = True

            except User.DoesNotExist:
                messages.error(request, "This email doesn't exist.")

        # === VERIFY CODE ===
        elif 'verify_code' in request.POST:
            input_code = request.POST.get('reset_code')
            session_code = request.session.get('reset_code')
            email_value = request.session.get('reset_email', '')

            if input_code == session_code:
                messages.success(request, "Code verified successfully! You can now reset your password.")
                return redirect('verify_reset_code')
            else:
                messages.error(request, "Invalid or incorrect code.")
                show_code_container = True

        # === RESEND CODE ===
        elif 'resend_code' in request.POST:
            email = request.session.get('reset_email')
            email_value = email

            if email:
                code = random.randint(100000, 999999)
                request.session['reset_code'] = str(code)

                # Reuse same HTML layout for the resend email
                subject = "🔐 TrailLend Password Reset Code (Resent)"
                from_email = settings.DEFAULT_FROM_EMAIL
                to = [email]

                html_content = f"""
                <!DOCTYPE html>
                <html lang="en">
                <head><meta charset="UTF-8" /></head>
                <body style="font-family:'Poppins',Arial,sans-serif; background-color:#f4f6f9; padding:30px;">
                  <table align="center" style="max-width:600px; background:#fff; border-radius:10px; padding:30px;">
                    <tr><td style="text-align:center;">
                      <img src="https://i.ibb.co/T2Hyfdd/TRAILLEND-ICON.png" width="70" alt="TrailLend" />
                      <h2 style="color:#1976D2;">TrailLend Password Reset (Resent)</h2>
                      <p style="color:#333;">Here is your new reset code:</p>
                      <div style="margin:20px 0;">
                        <span style="display:inline-block; background:#1976D2; color:#fff; font-size:26px; 
                                     letter-spacing:4px; padding:12px 22px; border-radius:8px; font-weight:bold;">
                          {code}
                        </span>
                      </div>
                      <p style="color:#555;">Enter this code in the TrailLend app to verify your identity.</p>
                      <p style="color:#999; font-size:12px;">If you didn’t request this, you can safely ignore this email.</p>
                    </td></tr>
                  </table>
                </body>
                </html>
                """

                text_content = strip_tags(html_content)

                email_message = EmailMultiAlternatives(subject, text_content, from_email, to)
                email_message.attach_alternative(html_content, "text/html")
                email_message.send()

                messages.success(request, "A new code has been sent to your email.")
                show_code_container = True
            else:
                messages.error(request, "No email session found. Please enter your email again.")

    # === RENDER TEMPLATE ===
    return render(request, "forgot_password.html", {
        'show_code_container': show_code_container,
        'email': email_value or request.session.get('reset_email', '')
    })


@csrf_exempt
def verify_reset_code(request):
    """
    Page for entering a new password after verifying the reset code.
    """
    email = request.session.get('reset_email')

    if not email:
        messages.error(request, "Session expired. Please enter your email again.")
        return redirect('forgot_password')

    if request.method == 'POST':
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')

        if not new_password or not confirm_password:
            messages.error(request, "Please fill in all fields.")
        elif new_password != confirm_password:
            messages.error(request, "Passwords do not match.")
        else:
            try:
                user = User.objects.get(email=email)
                user.set_password(new_password)
                user.save()

                # Clear session data
                request.session.pop('reset_email', None)
                request.session.pop('reset_code', None)

                # === Send confirmation email ===
                subject = "✅ Your TrailLend Password Has Been Changed"
                from_email = settings.DEFAULT_FROM_EMAIL
                to = [email]

                html_content = f"""
                <!DOCTYPE html>
                <html lang="en">
                <head>
                  <meta charset="UTF-8" />
                  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
                  <title>Password Changed - TrailLend</title>
                </head>
                <body style="margin:0; padding:0; background-color:#f4f6f9; font-family:'Poppins',Arial,sans-serif;">
                  <table align="center" border="0" cellpadding="0" cellspacing="0" width="100%" 
                         style="max-width:600px; background-color:#ffffff; border-radius:10px; overflow:hidden; 
                                box-shadow:0 4px 10px rgba(0,0,0,0.05); margin-top:40px;">
                    <tr>
                      <td style="background-color:#1976D2; text-align:center; padding:30px;">
                        <img src="https://i.ibb.co/T2Hyfdd/TRAILLEND-ICON.png" alt="TrailLend Logo" width="80" style="margin-bottom:10px;" />
                        <h1 style="color:#fff; font-size:22px; margin:0;">TrailLend</h1>
                        <p style="color:#cce6ff; font-size:13px; margin:4px 0 0;">Empowering the Community Together 🌿</p>
                      </td>
                    </tr>
                    <tr>
                      <td style="padding:30px;">
                        <h2 style="color:#1976D2; font-size:20px; margin-top:0;">Password Changed Successfully</h2>
                        <p style="color:#333; font-size:15px;">Hello {user.first_name or user.username},</p>
                        <p style="color:#333; font-size:15px;">
                          This is a confirmation that your password for your TrailLend account 
                          (<strong>{email}</strong>) has been successfully changed.
                        </p>

                        <p style="color:#555; font-size:14px; margin-top:20px;">
                          If you did not make this change, please contact your Barangay General Services Office 
                          immediately to secure your account.
                        </p>

                        <div style="text-align:center; margin:30px 0;">
                          <a href="https://traillend.com/login" 
                             style="display:inline-block; background:#1976D2; color:#fff; 
                                    padding:12px 25px; border-radius:6px; font-weight:bold; text-decoration:none;">
                            Go to TrailLend
                          </a>
                        </div>

                        <hr style="border:none; border-top:1px solid #eee; margin:30px 0;" />

                        <p style="color:#999; font-size:12px; text-align:center;">
                          This email was sent by <strong>TrailLend</strong> • Barangay General Services Office<br/>
                          Please do not reply directly to this message.
                        </p>
                      </td>
                    </tr>
                  </table>
                </body>
                </html>
                """

                text_content = strip_tags(html_content)
                email_message = EmailMultiAlternatives(subject, text_content, from_email, to)
                email_message.attach_alternative(html_content, "text/html")
                email_message.send()

                # Show success message on UI
                messages.success(request, "Your password has been successfully changed.")
                return render(request, "verify_reset_code.html")

            except User.DoesNotExist:
                messages.error(request, "User not found. Please try again.")

    return render(request, "verify_reset_code.html", {"email": email})


@csrf_exempt
def me_borrower(request):
    if not request.user.is_authenticated:
        return JsonResponse({"error": "Unauthorized"}, status=401)

    try:
        borrower = UserBorrower.objects.get(user=request.user)
        return JsonResponse({
            "user_id": borrower.id,
            "full_name": borrower.full_name,
            "contact_number": borrower.contact_number,
            "address": borrower.address,
            "late_count": borrower.late_count,
            "borrower_status": borrower.borrower_status,
        })
    except UserBorrower.DoesNotExist:
        return JsonResponse({"error": "Borrower profile not found"}, status=404)

@api_view(["GET"])
@permission_classes([IsAuthenticated])
def borrower_late_history(request):
    user = request.user
    try:
        borrower = user.userborrower
    except:
        return JsonResponse({"error": "Borrower not found"}, status=404)

    # Get all late returns
    late_feedback = Feedback.objects.filter(
        userborrower=borrower,
        return_status="Late"
    ).select_related("reservation__item")

    history = []
    for fb in late_feedback:
        reservation = fb.reservation
        item = reservation.item if reservation else None

        history.append({
            "reservation_id": reservation.id if reservation else None,
            "item_name": item.name if item else "Unknown Item",
            "date_borrowed": str(reservation.date_borrowed) if reservation else None,
            "date_return": str(reservation.date_return) if reservation else None,
            "feedback_date": fb.created_at.strftime("%Y-%m-%d %H:%M"),
        })

    data = {
        "full_name": borrower.full_name,
        "late_count": borrower.late_count,
        "borrower_status": borrower.borrower_status,
        "late_history": history,
    }

    return JsonResponse(data, safe=False)


def total_admin_borrow_for_date(item, target_date):
    qs = AdminBorrow.objects.filter(
        item=item,
        date__lte=target_date,
        return_date__gte=target_date
    ).aggregate(total=Sum("quantity"))
    return qs["total"] or 0


def total_reservation_qty_for_date(item, target_date):
    qs = Reservation.objects.filter(
        item=item,
        status__in=["pending", "approved", "in use"],
        date_borrowed__lte=target_date,
        date_return__gte=target_date,
    ).aggregate(total=Sum("quantity"))
    return qs["total"] or 0

def compute_daily_availability(item, target_date):
    total = item.qty
    reserved = total_reservation_qty_for_date(item, target_date)
    admin_used = total_admin_borrow_for_date(item, target_date)

    available = max(total - reserved - admin_used, 0)

    return {
        "total": total,
        "reserved": reserved,
        "admin_borrowed": admin_used,
        "available": available
    }

@api_view(["POST"])
@permission_classes([AllowAny])
def create_admin_borrow(request, item_id):
    from datetime import datetime, timedelta

    try:
        item = Item.objects.get(item_id=item_id)
    except Item.DoesNotExist:
        return Response({"error": "Item not found"}, status=404)

    data = request.data
    try:
        start_date = datetime.strptime(data["date"], "%Y-%m-%d").date()
        return_date = datetime.strptime(data["return_date"], "%Y-%m-%d").date()
    except:
        return Response({"error": "Invalid date format"}, status=400)

    qty = int(data.get("quantity", 0))
    if qty <= 0:
        return Response({"error": "Invalid quantity"}, status=400)

    # ------------- CHECK AVAILABILITY FOR EACH DAY -------------
    current = start_date
    while current <= return_date:
        reserved = Reservation.objects.filter(
            item=item,
            status__in=["pending", "approved", "in use"],
            date_borrowed__lte=current,
            date_return__gte=current
        ).aggregate(total=Sum("quantity"))["total"] or 0

        admin_used = AdminBorrow.objects.filter(
            item=item,
            status="In Use",
            date__lte=current,
            return_date__gte=current
        ).aggregate(total=Sum("quantity"))["total"] or 0

        available = item.qty - reserved - admin_used

        if available < qty:
            return Response({
                "error": f"Not enough available items on {current}. Only {available} left."
            }, status=400)

        current += timedelta(days=1)

    # ------------- SAVE ADMIN BORROW -------------
    ab = AdminBorrow.objects.create(
        item=item,
        date=start_date,
        return_date=return_date,
        quantity=qty,
        borrower_name=data["borrower_name"],
        contact_number=data["contact_number"],
        address=data.get("address", ""),
        purpose=data.get("purpose", ""),
        delivered_by=data["delivered_by"],
        status="In Use",
    )

    return Response({"message": "Admin borrow recorded", "transaction_id": ab.transaction_id})



@api_view(["PUT"])
@permission_classes([AllowAny])
def update_admin_borrow(request, pk):
    try:
        ab = AdminBorrow.objects.get(pk=pk)
        item = ab.item

        new_qty = int(request.data.get("quantity"))
        new_return = parse_date(request.data.get("return_date"))

        if new_qty < 1 or new_return < ab.date:
            return Response({"error": "Invalid input"}, status=400)

        # Validate entire new range
        current = ab.date
        while current <= new_return:
            avail = compute_daily_availability(item, current)
            # remove the old quantity first
            avail["available"] += ab.quantity

            if new_qty > avail["available"]:
                return Response({
                    "error": "Not enough availability",
                    "date": current.isoformat(),
                    "available": avail["available"]
                }, status=409)
            current += timedelta(days=1)

        ab.quantity = new_qty
        ab.return_date = new_return
        ab.save()

        return Response({"success": True})

    except AdminBorrow.DoesNotExist:
        return Response({"error": "Not found"}, status=404)

@api_view(["DELETE"])
@permission_classes([AllowAny])
def delete_admin_borrow(request, pk):
    try:
        ab = AdminBorrow.objects.get(pk=pk)
        ab.delete()
        return Response({"success": True})
    except AdminBorrow.DoesNotExist:
        return Response({"error": "Not found"}, status=404)

@api_view(["GET"])
@permission_classes([AllowAny])
def admin_borrow_list(request, item_id):
    date = request.GET.get("date")

    try:
        item = Item.objects.get(item_id=item_id)
    except Item.DoesNotExist:
        return Response([], status=200)

    # ALWAYS return all records — do NOT filter by date
    borrows = AdminBorrow.objects.filter(item=item).order_by("-id")

    output = []
    for ab in borrows:
        output.append({
            "id": ab.id,
            "transaction_id": ab.transaction_id,
            "borrower_name": ab.borrower_name,
            "contact_number": ab.contact_number,
            "quantity": ab.quantity,
            "return_date": str(ab.return_date),
            "status": ab.status,
            "delivered_by": ab.delivered_by,
        })

    return Response(output, status=200)


@csrf_exempt
def admin_borrow_create(request, item_id):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    try:
        data = json.loads(request.body)
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    # Extract data
    date = data.get("date")
    return_date = data.get("return_date")
    quantity = data.get("quantity")
    borrower_name = data.get("borrower_name")
    contact_number = data.get("contact_number")
    address = data.get("address", "")
    purpose = data.get("purpose", "")
    delivered_by = data.get("delivered_by", "")

    # Create direct borrow entry
    from core.models import Item, AdminBorrow

    try:
        item = Item.objects.get(pk=item_id)
    except Item.DoesNotExist:
        return JsonResponse({"error": "Item not found"}, status=404)

    ab = AdminBorrow.objects.create(
        item=item,
        date=date,
        return_date=return_date,
        quantity=int(quantity),
        borrower_name=borrower_name,
        contact_number=contact_number,
        address=address,
        purpose=purpose,
        delivered_by=delivered_by
    )

    return JsonResponse({
        "message": "Direct borrow successfully recorded.",
        "transaction_id": ab.transaction_id,
        "id": ab.id
    })

@csrf_exempt
@api_view(["POST"])
@permission_classes([AllowAny])
def return_admin_borrow(request, pk):
    try:
        ab = AdminBorrow.objects.get(pk=pk)
    except AdminBorrow.DoesNotExist:
        return Response({"error": "Not found"}, status=404)

    if ab.status == "Returned":
        return Response({"message": "Already returned"}, status=200)

    ab.status = "Returned"
    ab.return_date = date.today()  # ← UPDATE RETURN DATE
    ab.save()

    return Response({"message": "Marked as returned"}, status=200)

@csrf_exempt
def create_admin_borrow(request, item_id):
    data = json.loads(request.body)

    # VALIDATION (quantity, availability, etc. — I will send complete version later)

    ab = AdminBorrow.objects.create(
        item_id=item_id,
        date=data["date"],
        return_date=data["return_date"],
        quantity=data["quantity"],
        borrower_name=data["borrower_name"],
        contact_number=data["contact_number"],
        address=data.get("address", ""),
        purpose=data.get("purpose", ""),
        delivered_by=data["delivered_by"]
    )

    # ADD TO HISTORY OUTPUT FORMAT
    TransactionCounter.objects.create(
        transaction_id=ab.transaction_id,
        user_name=ab.borrower_name,
        item_name=ab.item.name,
        quantity=ab.quantity,
        contact=ab.contact_number,
        date_receive=ab.date,
        date_returned=None,
        delivered_by=ab.delivered_by,
        status="In Use",
    )

    return JsonResponse({"message": "Saved", "transaction_id": ab.transaction_id})


@login_required
def update_report_status(request, report_id):
    if request.method != "POST":
        return JsonResponse({"success": False, "message": "Invalid request method."}, status=400)

    action = request.POST.get("action")

    # ============================
    # FETCH REPORT
    # ============================
    try:
        report = DamageReport.objects.get(id=report_id)
    except DamageReport.DoesNotExist:
        return JsonResponse({"success": False, "message": "Report not found."}, status=404)

    borrower = report.reported_by
    item = report.item
    reservation = report.reservation  # may be None

    # ============================
    # STATUS TRANSITIONS
    # ============================
    if report.report_type == "Damage":
        if report.status == "Pending" and action == "review":
            report.status = "Reviewed"

        elif report.status == "Reviewed" and action == "resolve":
            report.status = "Resolved"

        else:
            return JsonResponse({"success": False, "message": "Invalid action for Damage report."}, status=400)

    elif report.report_type == "Loss":
        if report.status == "Pending" and action == "verify":
            report.status = "Verified"

        elif report.status == "Verified" and action == "settle":
            report.status = "Settled"

        else:
            return JsonResponse({"success": False, "message": "Invalid action for Loss report."}, status=400)

    # Save updated status first
    report.save()

    # ============================
    # APPLY CONSEQUENCES (only at FINAL stage)
    # ============================
    if report.status in ["Resolved", "Settled"]:

        # 1️⃣ Add violation points
        if report.report_type == "Damage":
            borrower.late_count += 1
        elif report.report_type == "Loss":
            borrower.late_count += 2

        # 2️⃣ Determine if borrower becomes “Bad”
        is_now_bad = False
        if borrower.late_count >= 3:
            borrower.borrower_status = "Bad"
            is_now_bad = True

        borrower.save()

        # ============================
        # 3️⃣ ITEM DEDUCTION — LOSS ONLY
        # ============================
        if report.report_type == "Loss" and not report.qty_deducted:

            lost_qty = report.quantity_affected

            # Deduct quantity safely
            item.qty = max(0, item.qty - lost_qty)
            item.save()

            # Mark as deducted
            report.qty_deducted = True
            report.save()

        # ============================
        # 4️⃣ Notifications
        # ============================

        transaction_id = reservation.transaction_id if reservation else None

        # Main report notification
        Notification.objects.create(
            user=borrower,
            reservation=reservation,
            title=f"{report.report_type} Report {report.status}",
            message=(
                f"Your {report.report_type.lower()} report for '{item.name}' is now {report.status}. "
                f"Violations: {borrower.late_count}/3."
            ),
            type="damage_report" if report.report_type == "Damage" else "loss_report",
        )

        # Borrower restricted
        if is_now_bad:
            Notification.objects.create(
                user=borrower,
                reservation=reservation,
                title="Borrower Privileges Revoked",
                message=(
                    "You now have 3 violations. Your TrailLend privileges have been restricted.\n"
                    "Please visit the GSO for re-evaluation."
                ),
                type="restricted",
            )

        # Warning for 1 or 2 violations
        elif borrower.late_count in [1, 2]:
            Notification.objects.create(
                user=borrower,
                reservation=reservation,
                title="Violation Warning",
                message=(
                    f"You now have {borrower.late_count}/3 violations. "
                    "You will be restricted once you reach 3."
                ),
                type="violation_warning",
            )

    # ============================
    # DONE
    # ============================
    return JsonResponse({"success": True, "new_status": report.status})


# ----------------------------------------------------------
# SUGGEST ALTERNATIVE ITEMS (DATE RANGE)
# ----------------------------------------------------------
@api_view(["POST"])
@permission_classes([AllowAny])
def suggest_items(request):
    try:
        start_date = parse_date(request.data.get("start_date"))
        end_date = parse_date(request.data.get("end_date"))
        exclude_item = request.data.get("exclude_item_id")

        if not start_date or not end_date:
            return Response({"success": False, "message": "Missing dates"}, status=400)

        if end_date < start_date:
            return Response({"success": False, "message": "Invalid date range"}, status=400)

        suggestions = []

        for item in Item.objects.all():

            if str(item.item_id) == str(exclude_item):
                continue

            overlap = ReservationItem.objects.filter(
                item=item,
                reservation__status__in=["pending", "approved", "in use"],
                reservation__date_borrowed__lte=end_date,
                reservation__date_return__gte=start_date,
            ).aggregate(total=Sum("quantity"))

            reserved = overlap["total"] or 0
            available = max(item.qty - reserved, 0)

            if available > 0:
                suggestions.append({
                    "item_id": item.item_id,
                    "name": item.name,
                    "owner": item.owner,
                    "description": item.description,
                    "available_qty": available,
                    "image": request.build_absolute_uri(item.image.url) if item.image else None
                })

        return Response({"success": True, "suggestions": suggestions})

    except Exception as e:
        return Response({"success": False, "message": str(e)}, status=500)



def schedule_smart_alerts(reservation):
    borrower = reservation.userborrower

    # 1️⃣ RETURN REMINDER — 6pm the day before date_return
    return_day_before = reservation.date_return - timedelta(days=1)
    return_dt = make_aware(datetime.combine(return_day_before, time(18, 0)))

    Notification.objects.create(
        user=borrower,
        reservation=reservation,
        title="Return Reminder",
        message="This is a reminder to return your borrowed items tomorrow.",
        type="return_reminder",
        scheduled_at=return_dt,
    )

    # 2️⃣ CLAIM REMINDER — 6am on date_borrowed
    claim_dt = make_aware(datetime.combine(reservation.date_borrowed, time(6, 0)))

    Notification.objects.create(
        user=borrower,
        reservation=reservation,
        title="Claim Reminder",
        message="You may now claim your reserved items today.",
        type="claim_reminder",
        scheduled_at=claim_dt,
    )

    # 3️⃣ CLAIM WARNING — 1 hour after reservation
    one_hour_later = reservation.created_at + timedelta(hours=1)

    Notification.objects.create(
        user=borrower,
        reservation=reservation,
        title="Claiming Delay Warning",
        message="You still haven't claimed your reserved items. Please claim them as soon as possible.",
        type="warning_claim_delay",
        scheduled_at=one_hour_later,
    )

@api_view(["POST"])
@parser_classes([MultiPartParser, FormParser])
@permission_classes([IsAuthenticated])
def submit_letter(request):
    try:
        borrower = UserBorrower.objects.get(user=request.user)

        # Extract normal fields
        letter_text = request.data.get("letter_text", "")
        date_borrowed = request.data.get("date_borrowed")
        date_return = request.data.get("date_return")
        priority = request.data.get("priority", "Low")
        priority_detail = request.data.get("priority_detail", "")
        contact = borrower.contact_number

        # Required
        if not date_borrowed or not date_return:
            return Response({"status": "error", "message": "Date range required"}, status=400)

        # Extract file uploads
        id_image = request.FILES.get("id_card_image")
        signature_img = request.FILES.get("signature_image")

        # Create reservation
        reservation = Reservation.objects.create(
            userborrower=borrower,
            letter_text=letter_text,
            date_borrowed=date_borrowed,
            date_return=date_return,
            priority=priority,
            priority_detail=priority_detail,
            contact=contact,
            valid_id_image=id_image,
            signature_image=signature_img,
            status="pending",

            # NEW SYSTEM FIELDS
            borrow_date=date_borrowed,               # System tracks start
            expected_return_date=date_return         # System tracks expected return
        )

        # Parse JSON items
        items_json = request.data.get("items")
        if not items_json:
            reservation.delete()
            return Response({"status": "error", "message": "Items list is required"}, status=400)

        try:
            items = json.loads(items_json)
        except:
            reservation.delete()
            return Response({"status": "error", "message": "Invalid items JSON"}, status=400)

        for it in items:
            item_obj = Item.objects.get(item_id=it["id"])
            qty = int(it["quantity"])

            ReservationItem.objects.create(
                reservation=reservation,
                item=item_obj,
                item_name=item_obj.name,
                quantity=qty
            )

# ⭐ Compute and store PRIORITY SCORE (AFTER items + created_at exist)
            reservation.refresh_from_db()  
            items = reservation.items.all()

            if items.exists():
                reservation.priority_score = compute_priority(reservation)
            else:
                reservation.priority_score = 0

            reservation.save()



        return Response({
            "status": "success",
            "transaction_id": reservation.transaction_id,
            "priority_score": reservation.priority_score
        }, status=201)

    except Exception as e:
        return Response({"status": "error", "message": str(e)}, status=500)
